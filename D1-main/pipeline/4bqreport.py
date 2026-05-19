"""
IPD Patent Analysis — Per-Drug Word Report Generator (Gemini-powered)
=====================================================================
Reads data from Google BigQuery tables:
  - shortlisted_secondary_patents_table  →  replaces "Shortlisted" sheet
  - arbitrage_summary_table              →  replaces "Arbitrage Summary" sheet

All BQ columns are snake_case and are mapped back to the original
Title Case names used throughout the report generation logic.

API key read from .env  →  GEMINI_API_KEY
Model: gemini-2.5-flash
"""

import os
import sys
import io
import re
import time
import tempfile
from datetime import date

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = lambda **kw: None  # Not needed on Cloud Run
import google.generativeai as genai
from google.cloud import bigquery
from google.oauth2 import service_account

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════
#  BIGQUERY CONFIG
# ═══════════════════════════════════════════════════════════════════════════
BQ_PROJECT_ID  = "cognito-prod-394707"
BQ_DATASET_ID  = "cognito_prod_datamart"
BQ_TABLE_ID    = "Master_LOE"          # not used directly; kept for reference
CREDENTIALS_PATH = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "")
BQ_LOCATION    = "asia-south1"

# Table names
SHORTLISTED_TABLE = "shortlisted_secondary_patents_table"
ARBITRAGE_TABLE   = "arbitrage_summary_table"

# ── snake_case  →  Title Case column mappings ─────────────────────────────
SHORTLISTED_COL_MAP = {
    "drug_name":                    "Drug Name",
    "jurisdiction":                 "Jurisdiction",
    "patent_number":                "Patent Number",
    "step_1_claim_category":        "Step 1 Claim Category",
    "adjusted_expiry_with_pte":     "Adjusted Expiry (with PTE)",
    "expiry_gap_years":             "Expiry Gap (Years)",
    "pte_status":                   "PTE Status",
    "pte_months_granted":           "PTE Months (Granted)",
}

ARBITRAGE_COL_MAP = {
    "drug_name":                    "Drug Name",
    "jurisdiction":                 "Jurisdiction",
    "dimension_iv_score":           "Dimension IV Score",
    "dimension_iv_rating":          "Dimension IV Rating",
    "product_loe_year":             "Product LOE (Year)",
    "gap_vs_us_years":              "Gap vs US (Years)",
    "gap_vs_longest_loe_years":     "Gap vs Longest LOE (Years)",
    "key_protection_gap":           "Key Protection Gap",
    "arbitrage_score":              "Arbitrage Score",
    "arbitrage_signal":             "Arbitrage Signal",
    "rationale":                    "Rationale",
    "created_at":                   "Created At",
    "updated_at":                   "Updated At",
}



# ═══════════════════════════════════════════════════════════════════════════
#  BIGQUERY SCHEMA MIGRATION  —  create columns if absent
# ═══════════════════════════════════════════════════════════════════════════
def _ensure_arbitrage_columns(client: bigquery.Client) -> None:
    """
    Add 'rationale' (STRING), 'created_at' (TIMESTAMP), and 'updated_at'
    (TIMESTAMP) columns to arbitrage_summary_table if they do not already
    exist, then populate them appropriately.

    - created_at  : set once (first run) and never overwritten.
    - updated_at  : set to CURRENT_TIMESTAMP() on every run.
    """
    full_table = f"{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{ARBITRAGE_TABLE}"

    # Fetch current schema
    table = client.get_table(full_table)
    existing = {f.name for f in table.schema}

    ddl_statements = []

    if "rationale" not in existing:
        ddl_statements.append(
            f"ALTER TABLE `{full_table}` ADD COLUMN IF NOT EXISTS rationale STRING"
        )

    if "created_at" not in existing:
        ddl_statements.append(
            f"ALTER TABLE `{full_table}` ADD COLUMN IF NOT EXISTS created_at TIMESTAMP"
        )

    if "updated_at" not in existing:
        ddl_statements.append(
            f"ALTER TABLE `{full_table}` ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP"
        )

    for stmt in ddl_statements:
        print(f"  [BQ DDL] {stmt}")
        client.query(stmt).result()
        print(f"           ✓ Done")

    # Populate rationale where NULL
    update_rationale = f"""
        UPDATE `{full_table}`
        SET rationale = CONCAT(
            'Arbitrage Score is ', CAST(arbitrage_score AS STRING),
            ' because of ', CAST(arbitrage_signal AS STRING)
        )
        WHERE rationale IS NULL
    """
    print("  [BQ] Populating rationale column …")
    client.query(update_rationale).result()
    print("       ✓ rationale populated")

    # Populate created_at only where NULL (i.e. first time this row is seen)
    update_created_at = f"""
        UPDATE `{full_table}`
        SET created_at = CURRENT_TIMESTAMP()
        WHERE created_at IS NULL
    """
    print("  [BQ] Populating created_at column (first-run rows only) …")
    client.query(update_created_at).result()
    print("       ✓ created_at populated")

    # Always refresh updated_at to reflect the current run
    update_updated_at = f"""
        UPDATE `{full_table}`
        SET updated_at = CURRENT_TIMESTAMP()
        WHERE TRUE
    """
    print("  [BQ] Refreshing updated_at column (all rows) …")
    client.query(update_updated_at).result()
    print("       ✓ updated_at refreshed")


# ═══════════════════════════════════════════════════════════════════════════
#  BIGQUERY LOADER
# ═══════════════════════════════════════════════════════════════════════════
def _get_credentials():
    """Get credentials: use service account file if available, else default (Cloud Run)."""
    if CREDENTIALS_PATH and os.path.exists(CREDENTIALS_PATH):
        return service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)
    return None  # Use ADC (Application Default Credentials)

def _get_bq_client() -> bigquery.Client:
    """Return an authenticated BigQuery client."""
    credentials = _get_credentials()
    return bigquery.Client(
        project=BQ_PROJECT_ID,
        credentials=credentials,
        location=BQ_LOCATION,
    )


def _load_table(client: bigquery.Client, table_name: str) -> pd.DataFrame:
    """Load an entire BQ table into a DataFrame (all columns as strings)."""
    full_ref = f"`{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{table_name}`"
    query    = f"SELECT * FROM {full_ref}"
    print(f"  [BQ] Querying {full_ref} …")
    df = client.query(query).to_dataframe().astype(str)
    # Replace literal "None" / "nan" strings with proper NaN
    df.replace({"None": pd.NA, "nan": pd.NA, "<NA>": pd.NA}, inplace=True)
    print(f"       → {len(df)} rows, columns: {list(df.columns)}")
    return df


def _rename_columns(df: pd.DataFrame, col_map: dict) -> pd.DataFrame:
    """
    Rename snake_case BQ columns to Title Case display names.
    Only renames columns that exist; unknown columns are left as-is.
    """
    existing_map = {k: v for k, v in col_map.items() if k in df.columns}
    df = df.rename(columns=existing_map)

    # Warn about expected columns that were missing in the table
    missing = [v for k, v in col_map.items() if k not in existing_map]
    if missing:
        print(f"  ⚠  The following expected columns were NOT found in BQ: {missing}")
    return df


def load_data_from_bigquery() -> tuple[pd.DataFrame, pd.DataFrame]:
    """
    Returns (shortlisted_df, arbitrage_df) with Title Case columns,
    ready for the report generation functions.
    """
    client = _get_bq_client()

    # Ensure rationale + created_at + updated_at columns exist in BQ
    print("  Ensuring arbitrage table schema …")
    _ensure_arbitrage_columns(client)

    # Shortlisted
    sl_raw = _load_table(client, SHORTLISTED_TABLE)
    sl_df  = _rename_columns(sl_raw, SHORTLISTED_COL_MAP)
    sl_df.columns = sl_df.columns.str.strip()

    # Arbitrage Summary
    arb_raw = _load_table(client, ARBITRAGE_TABLE)
    arb_df  = _rename_columns(arb_raw, ARBITRAGE_COL_MAP)
    arb_df.columns = arb_df.columns.str.strip()

    print(f"\n  ✓ Shortlisted  : {len(sl_df)} rows | "
          f"{sl_df['Drug Name'].nunique() if 'Drug Name' in sl_df.columns else '?'} drug(s)")
    print(f"  ✓ Arbitrage    : {len(arb_df)} rows")

    return sl_df, arb_df


# ═══════════════════════════════════════════════════════════════════════════
#  GEMINI SETUP
# ═══════════════════════════════════════════════════════════════════════════
load_dotenv(override=True)
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
gemini = genai.GenerativeModel("gemini-2.5-flash")


def _call_gemini(prompt: str, retries: int = 3, backoff: float = 2.5) -> str:
    for attempt in range(retries):
        try:
            resp = gemini.generate_content(prompt)
            return resp.text.strip()
        except Exception as e:
            if attempt < retries - 1:
                wait = backoff * (2 ** attempt)
                print(f"    ⚠ Gemini error ({e}), retrying in {wait}s …")
                time.sleep(wait)
            else:
                print(f"    ✗ Gemini failed after {retries} attempts: {e}")
                return "Narrative unavailable."


# ═══════════════════════════════════════════════════════════════════════════
#  COLOUR PALETTE
# ═══════════════════════════════════════════════════════════════════════════
BRAND_BLUE     = RGBColor(0x2F, 0x55, 0x97)
ACCENT_BLUE    = RGBColor(0x44, 0x72, 0xC4)
LIGHT_BLUE_HEX = "DCE6F1"
HEADER_HEX     = "2F5597"
WHITE_HEX      = "FFFFFF"

SCORE_HEX = {5: "2ECC71", 4: "27AE60", 3: "F39C12", 2: "E67E22", 1: "E74C3C"}
SIGNAL_HEX = {
    "Immediate opportunity": "2ECC71",
    "Strong arbitrage":      "27AE60",
    "Meaningful arbitrage":  "F39C12",
    "Limited arbitrage":     "E67E22",
    "No arbitrage":          "E74C3C",
    "Reference market":      "3498DB",
}
RATING_HEX = {
    "Very Strong Geographic Arbitrage Opportunity": "2ECC71",
    "Strong Geographic Arbitrage Opportunity":      "27AE60",
    "Moderate Geographic Arbitrage Opportunity":    "F39C12",
    "Limited Geographic Arbitrage Opportunity":     "E67E22",
    "FAIL — No viable geographic arbitrage":        "E74C3C",
}
JUR_FULL = {
    "CN": "China", "IN": "India", "BR": "Brazil",
    "AU": "Australia", "RU": "Russia", "US": "United States",
    "JP": "Japan", "KR": "South Korea", "TW": "Taiwan",
    "CA": "Canada", "MX": "Mexico",
}


# ═══════════════════════════════════════════════════════════════════════════
#  XML / STYLE HELPERS
# ═══════════════════════════════════════════════════════════════════════════
BORDER_OPTS = {"val": "single", "sz": 4, "color": "CCCCCC"}


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, opts in kwargs.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   opts.get("val",   "single"))
        el.set(qn("w:sz"),    str(opts.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), opts.get("color", "CCCCCC"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _page_break(doc: Document):
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_para_border_bottom(para, color="2F5597", sz=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _make_header_cell(cell, text):
    _set_cell_bg(cell, HEADER_HEX)
    _set_cell_border(cell, top=BORDER_OPTS, bottom=BORDER_OPTS,
                     left=BORDER_OPTS, right=BORDER_OPTS)
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _make_data_cell(cell, text, bg=WHITE_HEX, font_size=9):
    _set_cell_bg(cell, bg)
    _set_cell_border(cell, top=BORDER_OPTS, bottom=BORDER_OPTS,
                     left=BORDER_OPTS, right=BORDER_OPTS)
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text) if text is not None else "—")
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _signal_cell(cell, signal_text: str):
    hex_c = SIGNAL_HEX.get(signal_text, "FFFFFF")
    _set_cell_bg(cell, hex_c)
    _set_cell_border(cell, top=BORDER_OPTS, bottom=BORDER_OPTS,
                     left=BORDER_OPTS, right=BORDER_OPTS)
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(signal_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _score_cell(cell, score_val):
    try:
        hex_c = SCORE_HEX.get(int(score_val), "FFFFFF")
    except (ValueError, TypeError):
        hex_c = "BDC3C7"
    _set_cell_bg(cell, hex_c)
    _set_cell_border(cell, top=BORDER_OPTS, bottom=BORDER_OPTS,
                     left=BORDER_OPTS, right=BORDER_OPTS)
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(score_val))
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fig_to_docx(doc, fig, width_inches=5.5):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    doc.add_picture(buf, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
#  DOCUMENT SETUP
# ═══════════════════════════════════════════════════════════════════════════
def _setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.0)

    for name, size, color in [
        ("Heading 1", 15, BRAND_BLUE),
        ("Heading 2", 12, ACCENT_BLUE),
        ("Heading 3", 11, BRAND_BLUE),
    ]:
        try:
            s = doc.styles[name]
            s.font.size      = Pt(size)
            s.font.color.rgb = color
            s.font.bold      = True
            s.font.name      = "Calibri"
        except KeyError:
            pass

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    return doc


# ═══════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
def _cover_page(doc, drug_name: str):
    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(drug_name)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Calibri"

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Global Launch Sequencing & Geographic Arbitrage Analysis")
    run2.italic = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = ACCENT_BLUE
    run2.font.name = "Calibri"

    doc.add_paragraph()
    _set_para_border_bottom(doc.add_paragraph(), sz=16)
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p3.add_run("Report Date:  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Calibri"
    r2 = p3.add_run(date.today().strftime("%d %B %Y"))
    r2.font.size = Pt(12)
    r2.font.name = "Calibri"
    r2.font.color.rgb = BRAND_BLUE

    _page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════
#  LLM PROMPT
# ═══════════════════════════════════════════════════════════════════════════
def _build_drug_narrative_prompt(drug, drug_sl, drug_arb) -> str:
    exp_lines = []
    for _, r in drug_sl.iterrows():
        jur    = JUR_FULL.get(str(r.get("Jurisdiction", "")), str(r.get("Jurisdiction", "")))
        cat    = r.get("Step 1 Claim Category", "")
        expiry = str(r.get("Adjusted Expiry (with PTE)", "N/A"))
        gap    = str(r.get("Expiry Gap (Years)", "N/A"))
        pte_st = r.get("PTE Status", "")
        pte_m  = r.get("PTE Months (Granted)", "0")
        pat_no = r.get("Patent Number", "N/A")
        exp_lines.append(
            f"  {jur} | Patent: {pat_no} | {cat} | Expiry: {expiry} | Gap: {gap} yrs"
            f" | PTE: {pte_st} ({pte_m} months)"
        )

    arb_lines = []
    dim4_score  = "N/A"
    dim4_rating = "N/A"
    if not drug_arb.empty:
        dim4_score  = drug_arb["Dimension IV Score"].iloc[0]
        dim4_rating = drug_arb["Dimension IV Rating"].iloc[0]
        for _, r in drug_arb.iterrows():
            jur         = JUR_FULL.get(str(r.get("Jurisdiction", "")), str(r.get("Jurisdiction", "")))
            loe         = r.get("Product LOE (Year)", "N/A")
            gap_us      = r.get("Gap vs US (Years)", "N/A")
            gap_longest = r.get("Gap vs Longest LOE (Years)", "N/A")
            signal      = r.get("Arbitrage Signal", "N/A")
            gap_kp      = r.get("Key Protection Gap", "N/A")
            arb_lines.append(
                f"  {jur} | LOE: {loe} | Gap vs US: {gap_us} yrs"
                f" | Gap vs Longest LOE: {gap_longest} yrs"
                f" | Signal: {signal} | Protection gap: {gap_kp}"
            )

    jurs_in_data = set()
    for _, r in drug_sl.iterrows():
        jurs_in_data.add(str(r.get("Jurisdiction", "")))
    if not drug_arb.empty:
        for _, r in drug_arb.iterrows():
            jurs_in_data.add(str(r.get("Jurisdiction", "")))

    all_target_jurs   = {"CN", "IN", "BR", "AU", "RU", "US", "CA", "JP", "MX", "TW", "KR"}
    jurs_present      = jurs_in_data & all_target_jurs
    jurs_full_names   = ", ".join(sorted(JUR_FULL.get(j, j) for j in all_target_jurs))
    jurs_analysed_names = ", ".join(sorted(JUR_FULL.get(j, j) for j in jurs_present))

    return f"""
You are a senior pharmaceutical patent analyst writing a drug-level patent
assessment for an internal IP strategy report.

DRUG: {drug}

PATENT EXPIRY DATA (one row per jurisdiction):
{chr(10).join(exp_lines) if exp_lines else '  No data.'}

GEOGRAPHIC ARBITRAGE DATA (one row per jurisdiction):
{chr(10).join(arb_lines) if arb_lines else '  No data.'}

OVERALL DIMENSION IV SCORE: {dim4_score} — {dim4_rating}

TARGET JURISDICTIONS CONSIDERED: {jurs_full_names}
JURISDICTIONS WITH BLOCKING PATENTS ANALYSED: {jurs_analysed_names}

TASK:
Write a detailed, analytical drug-level narrative (220–280 words) structured
into exactly these three sections. Use the exact plain-text section headers shown
below on their own line (no markdown, no asterisks, no bullets):

PATENT LANDSCAPE OVERVIEW
Write 7–8 sentences summarising the breadth and strength of patent protection
for this drug across all jurisdictions, including the range of expiry years
and whether PTE extensions have been granted.
Do NOT give analysis of US. It is only for reference. Analyse other jurisdictions
based on the US information. Mention the Patent Numbers.
Mention that patents from {jurs_full_names} were considered, but only the blocking
ones were analysed. Jurisdictions not present in the analysis data lacked blocking
patents — do NOT say "no data is available", instead say they were considered but
not included as they lacked blocking patents.
Highlight any jurisdictions with particularly strong or weak patent protection,
and note the overall Dimension IV Score and Rating for the drug.

GEOGRAPHIC ARBITRAGE ANALYSIS
Write 10–12 sentences analysing where the strongest and weakest arbitrage
opportunities exist. Reference specific jurisdictions, LOE years, gap vs US,
gap vs longest LOE, and arbitrage signal ratings. Explain the strategic
significance of the "Gap vs Longest LOE" metric — it measures how many years
earlier a jurisdiction loses exclusivity compared to the jurisdiction with the
longest protection.

KEY PROTECTION GAPS
Write 7–8 sentences identifying any missing claim categories or jurisdictions
with weak or absent protection, and explain what this means for a potential
generic or biosimilar entrant.

Rules:
- Do NOT include any recommendations or strategic advice sections.
- Formal, precise tone. No filler phrases.
- Do NOT give Analysis for US patents. Compare other Patents with the US information.
- Use full jurisdiction names, not abbreviations.
- Do NOT invent any numbers or dates — use only the data provided above.
- Plain text only. No markdown, no bold, no asterisks, no bullet points.
""".strip()


# ═══════════════════════════════════════════════════════════════════════════
#  CHART
# ═══════════════════════════════════════════════════════════════════════════
def _chart_drug_loe(drug: str, drug_arb: pd.DataFrame):
    df = drug_arb[drug_arb["Jurisdiction"] != "US"].copy()
    df["LOE"] = pd.to_numeric(df["Product LOE (Year)"], errors="coerce")
    df = df.dropna(subset=["LOE"]).sort_values("LOE")
    if df.empty:
        return None

    today  = date.today().year
    labels = [JUR_FULL.get(j, j) for j in df["Jurisdiction"]]
    colors = ["#" + SIGNAL_HEX.get(s, "BDC3C7") for s in df["Arbitrage Signal"]]

    fig, ax = plt.subplots(figsize=(5.5, max(2.5, 0.55 * len(df) + 1)))
    bars = ax.barh(labels, df["LOE"] - today, left=today,
                   color=colors, edgecolor="white", linewidth=0.8)
    ax.axvline(x=today, color="#E74C3C", linestyle="--", linewidth=1.5, label="Today")
    ax.set_xlabel("Year", fontsize=9)
    ax.set_title(f"{drug} — Loss of Exclusivity by Jurisdiction",
                 fontsize=10, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for bar, loe in zip(bars, df["LOE"]):
        ax.text(bar.get_x() + bar.get_width() - 0.2,
                bar.get_y() + bar.get_height() / 2,
                str(int(loe)), va="center", ha="right",
                fontsize=8, fontweight="bold", color="white")

    patches = [mpatches.Patch(color="#" + v, label=k) for k, v in SIGNAL_HEX.items()]
    ax.legend(handles=patches, fontsize=7, loc="lower right", framealpha=0.8, ncol=1)
    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════
#  RENDER LLM NARRATIVE
# ═══════════════════════════════════════════════════════════════════════════
NARRATIVE_HEADERS = [
    "PATENT LANDSCAPE OVERVIEW",
    "GEOGRAPHIC ARBITRAGE ANALYSIS",
    "KEY PROTECTION GAPS",
]


def _render_narrative(doc, narrative: str):
    pattern = "(" + "|".join(re.escape(h) for h in NARRATIVE_HEADERS) + ")"
    parts   = re.split(pattern, narrative)

    for part in parts:
        part = part.strip()
        if not part:
            continue
        if "STRATEGIC RECOMMENDATIONS" in part.upper() or "RECOMMENDATION" in part.upper():
            continue
        if part in NARRATIVE_HEADERS:
            p   = doc.add_paragraph()
            run = p.add_run(part.title())
            run.bold           = True
            run.font.size      = Pt(11)
            run.font.name      = "Calibri"
            run.font.color.rgb = BRAND_BLUE
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
        else:
            p   = doc.add_paragraph()
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════
#  PATENT EXPIRY SUMMARY
# ═══════════════════════════════════════════════════════════════════════════
def _patent_expiry_summary(doc, drug_sl: pd.DataFrame):
    p   = doc.add_paragraph()
    run = p.add_run("Patent Expiry Summary")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = BRAND_BLUE
    _set_para_border_bottom(p, sz=6)

    exp_cols = ["Patent Number", "Jurisdiction", "Step 1 Claim Category",
                "Adjusted Expiry (with PTE)", "Expiry Gap (Years)",
                "PTE Status", "PTE Months (Granted)"]
    exp_cols = [c for c in exp_cols if c in drug_sl.columns]
    exp_df   = drug_sl[exp_cols].copy()

    if "Adjusted Expiry (with PTE)" in exp_df.columns:
        exp_df["Adjusted Expiry (with PTE)"] = pd.to_datetime(
            exp_df["Adjusted Expiry (with PTE)"], errors="coerce"
        ).dt.strftime("%Y-%m-%d").fillna("Not found")

    etbl = doc.add_table(rows=1 + len(exp_df), cols=len(exp_cols))
    etbl.style     = "Table Grid"
    etbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(exp_cols):
        _make_header_cell(etbl.rows[0].cells[i], h)

    for r_idx, (_, row_data) in enumerate(exp_df.iterrows()):
        bg  = LIGHT_BLUE_HEX if r_idx % 2 == 0 else WHITE_HEX
        row = etbl.rows[r_idx + 1]
        for c_idx, col in enumerate(exp_cols):
            _make_data_cell(row.cells[c_idx], row_data[col], bg=bg)

    exp_widths = [0.8, 1.8, 1.4, 1.0, 1.2, 1.1][:len(exp_cols)]
    for ci, w in enumerate(exp_widths):
        for cell in etbl.columns[ci].cells:
            cell.width = Inches(w)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
#  PDF CONVERSION & GCS UPLOAD
# ═══════════════════════════════════════════════════════════════════════════
def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert a .docx file to PDF using docx2pdf.

    On Windows this uses Microsoft Word via COM automation (Word must be
    installed).  On macOS it uses the Word AppleScript bridge.  No external
    command-line tools are required.

    Install once with:
        pip install docx2pdf

    Returns the path to the generated PDF.
    """
    try:
        from docx2pdf import convert as _docx2pdf_convert
    except ImportError:
        raise RuntimeError(
            "docx2pdf is not installed.\n"
            "Run: pip install docx2pdf"
        )

    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    _docx2pdf_convert(docx_path, pdf_path)

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF not created at expected path: {pdf_path}")
    return pdf_path


def upload_to_gcs(local_pdf: str, drug_name: str,
                  bucket_name: str = "cognito-gcs") -> str:
    """Upload PDF to gs://{bucket_name}/Cognito_new/reports/{drug_name}/Loe_Report(Secondary_Market).pdf
    Returns the full GCS URI."""
    try:
        from google.cloud import storage
    except ImportError:
        raise RuntimeError(
            "google-cloud-storage not installed.\n"
            "Run: pip install google-cloud-storage --break-system-packages"
        )

    safe_drug = drug_name.replace("/", "_").replace(" ", "_")
    blob_name = f"Cognito_new/reports/{safe_drug}/Loe_Report(Secondary_Market).pdf"

    credentials = _get_credentials()
    client = storage.Client(project=BQ_PROJECT_ID, credentials=credentials)

    bucket = client.bucket(bucket_name)
    blob   = bucket.blob(blob_name)
    blob.upload_from_filename(local_pdf, content_type="application/pdf")

    gcs_uri = f"gs://{bucket_name}/{blob_name}"
    print(f"      ✓ Uploaded → {gcs_uri}")
    return gcs_uri


# ═══════════════════════════════════════════════════════════════════════════
#  BUILD SINGLE-DRUG REPORT
# ═══════════════════════════════════════════════════════════════════════════
def _build_drug_report(drug, drug_sl, drug_arb, output_dir):
    doc = _setup_document()

    # Cover
    _cover_page(doc, drug)

    # Dimension IV badge
    if not drug_arb.empty:
        dim4_score  = drug_arb["Dimension IV Score"].iloc[0]
        dim4_rating = drug_arb["Dimension IV Rating"].iloc[0]
        rating_hex  = RATING_HEX.get(str(dim4_rating), "BDC3C7")

        badge = doc.add_table(rows=1, cols=2)
        badge.alignment = WD_TABLE_ALIGNMENT.LEFT

        c0 = badge.rows[0].cells[0]
        _set_cell_bg(c0, HEADER_HEX)
        _set_cell_border(c0, top=BORDER_OPTS, bottom=BORDER_OPTS,
                         left=BORDER_OPTS, right=BORDER_OPTS)
        p   = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Dimension IV Score")
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10); c0.width = Inches(1.8)

        c1 = badge.rows[0].cells[1]
        _set_cell_bg(c1, rating_hex)
        _set_cell_border(c1, top=BORDER_OPTS, bottom=BORDER_OPTS,
                         left=BORDER_OPTS, right=BORDER_OPTS)
        p   = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{dim4_score}  —  {dim4_rating}")
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10); c1.width = Inches(4.5)

        doc.add_paragraph()

    # Geographic Arbitrage Map Table
    if not drug_arb.empty:
        p   = doc.add_paragraph()
        run = p.add_run("Geographic Arbitrage Map")
        run.bold = True; run.font.size = Pt(11)
        run.font.name = "Calibri"; run.font.color.rgb = BRAND_BLUE
        _set_para_border_bottom(p, sz=6)

        arb_cols = ["Jurisdiction", "Product LOE (Year)", "Gap vs US (Years)",
                    "Gap vs Longest LOE (Years)", "Key Protection Gap",
                    "Arbitrage Score", "Arbitrage Signal",
                    "Rationale", "Created At", "Updated At"]
        arb_cols = [c for c in arb_cols if c in drug_arb.columns]
        arb_show = drug_arb[arb_cols].copy()

        sig_idx   = arb_cols.index("Arbitrage Signal") if "Arbitrage Signal" in arb_cols else -1
        score_idx = arb_cols.index("Arbitrage Score")  if "Arbitrage Score"  in arb_cols else -1

        atbl = doc.add_table(rows=1 + len(arb_show), cols=len(arb_cols))
        atbl.style = "Table Grid"; atbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(arb_cols):
            _make_header_cell(atbl.rows[0].cells[i], h)
        for r_idx, (_, row_data) in enumerate(arb_show.iterrows()):
            bg  = LIGHT_BLUE_HEX if r_idx % 2 == 0 else WHITE_HEX
            row = atbl.rows[r_idx + 1]
            for c_idx, col in enumerate(arb_cols):
                cell = row.cells[c_idx]
                val  = str(row_data[col]) if row_data[col] is not None else "—"
                if c_idx == sig_idx:
                    _signal_cell(cell, val)
                elif c_idx == score_idx:
                    _score_cell(cell, val)
                else:
                    _make_data_cell(cell, val, bg=bg)
        arb_widths = [0.85, 0.9, 0.9, 1.0, 1.8, 0.8, 1.1, 2.2, 1.2, 1.2][:len(arb_cols)]
        for ci, w in enumerate(arb_widths):
            for cell in atbl.columns[ci].cells:
                cell.width = Inches(w)
        doc.add_paragraph()

        fig = _chart_drug_loe(drug, drug_arb)
        if fig:
            _fig_to_docx(doc, fig, width_inches=5.0)

    # LLM Narrative
    p   = doc.add_paragraph()
    run = p.add_run("Analysis")
    run.bold = True; run.font.size = Pt(11)
    run.font.name = "Calibri"; run.font.color.rgb = BRAND_BLUE
    _set_para_border_bottom(p, sz=6)

    prompt    = _build_drug_narrative_prompt(drug, drug_sl, drug_arb)
    narrative = _call_gemini(prompt)
    time.sleep(1)
    _render_narrative(doc, narrative)

    # Patent Expiry Summary (last)
    _page_break(doc)
    _patent_expiry_summary(doc, drug_sl)

    output_path = os.path.join(output_dir, "Loe_Report(Secondary_Market).docx")
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
#  TERMINAL HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def _get_output_dir() -> str:
    while True:
        path = input("\nEnter output DIRECTORY for per-drug reports: ").strip().strip('"').strip("'")
        if not path:
            print("  ✗ Path cannot be empty."); continue
        if not os.path.isdir(path):
            try:
                os.makedirs(path, exist_ok=True)
                print(f"  ✓ Created: {path}")
            except OSError as e:
                print(f"  ✗ Cannot create directory: {e}"); continue
        else:
            print(f"  ✓ Output dir: {path}")
        return path


# ═══════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════
def main():
    print("=" * 60)
    print("  IPD PER-DRUG REPORT GENERATOR  —  Gemini 2.5 Flash")
    print("  Data source: Google BigQuery")
    print("  Output     : gs://cognito-gcs/Cognito_new/reports/{drug_name}/Loe_Report(Secondary_Market).pdf")
    print("=" * 60)

    if not os.getenv("GEMINI_API_KEY"):
        print("  ✗ GEMINI_API_KEY not set — aborting.")
        sys.exit(1)

    # Load data from BigQuery
    print("\n  Loading data from BigQuery…")
    shortlisted, arb_df = load_data_from_bigquery()

    if shortlisted.empty:
        print("  ✗ Shortlisted table is empty — aborting.")
        sys.exit(1)

    if "Drug Name" not in shortlisted.columns:
        print("  ✗ 'drug_name' column not found in shortlisted table — check mapping.")
        sys.exit(1)

    # Generate one report per drug
    drugs = sorted(shortlisted["Drug Name"].dropna().unique())
    print(f"\n  Generating {len(drugs)} report(s)…\n")

    uploaded_uris = []
    with tempfile.TemporaryDirectory() as tmpdir:
        for idx, drug in enumerate(drugs, start=1):
            print(f"  [{idx}/{len(drugs)}] {drug}")

            drug_sl  = shortlisted[shortlisted["Drug Name"] == drug].copy()
            drug_arb = (
                arb_df[arb_df["Drug Name"] == drug].copy()
                if not arb_df.empty and "Drug Name" in arb_df.columns
                else pd.DataFrame()
            )

            # Build DOCX in a per-drug subfolder to avoid filename collisions
            drug_tmp = os.path.join(tmpdir, re.sub(r'[^\w\s-]', '', drug).strip().replace(' ', '_'))
            os.makedirs(drug_tmp, exist_ok=True)

            docx_path = _build_drug_report(drug, drug_sl, drug_arb, drug_tmp)
            print(f"      ✓ DOCX built")

            print(f"      Converting to PDF …", end=" ", flush=True)
            try:
                pdf_path = convert_docx_to_pdf(docx_path)
                print("done")
            except RuntimeError as exc:
                print(f"FAILED\n      ✗ {exc}")
                continue

            print(f"      Uploading to GCS …", end=" ", flush=True)
            try:
                uri = upload_to_gcs(pdf_path, drug)
                uploaded_uris.append(uri)
            except Exception as exc:
                print(f"FAILED\n      ✗ {exc}")

    print(f"\n{'='*60}")
    print(f"  ✓ Done — {len(uploaded_uris)}/{len(drugs)} PDF(s) uploaded")
    for uri in uploaded_uris:
        print(f"    {uri}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
