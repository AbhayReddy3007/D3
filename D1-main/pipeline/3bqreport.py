"""
Patent Thicket Report Generator — BigQuery Edition
====================================================
Reads patent thicket data from two BigQuery tables written by
patent_thicket_analysis.py and generates a polished Word document
with ~2 pages per drug, using Gemini to produce narrative prose.

BigQuery tables (written by patent_thicket_analysis.py):
  Patent_Thicket_Score_Table   — one row per drug+jurisdiction + avg row
  Circumvention_Table          — one row per design-around strategy

Exact BQ column names are taken directly from write_score_to_bq()
and write_circumvention_to_bq() in patent_thicket_analysis.py.

Usage:
    python generate_patent_report_bq.py [output_docx] [--api-key KEY]

    # Override any config value at runtime:
    python generate_patent_report_bq.py report.docx \\
        --project cognito-prod-394707 \\
        --dataset cognito_prod_datamart \\
        --key-file "C:\\path\\to\\key.json" \\
        --location asia-south1

    # Create a .env file with:
    #   GEMINI_API_KEY=your-gemini-key
    #   GCS_CREDENTIALS=/path/to/service-account.json

Requirements:
    pip install pandas google-cloud-bigquery db-dtypes google-auth \\
                python-docx google-generativeai python-dotenv \\
                mammoth xhtml2pdf --break-system-packages
"""

import os
import re
import sys
import json
import shutil
import argparse
import tempfile
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

try:
    # google.generativeai is deprecated. Use google.genai (new SDK) only.
    from google import genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("[WARN] google-genai not installed.")

# ── Load .env ─────────────────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv(override=True)
except ImportError:
    pass  # python-dotenv not installed; fall back to environment variables

# ── BigQuery / GCP Config ─────────────────────────────────────────────────────
PROJECT_ID       = os.environ.get("PROJECT_ID",       "cognito-prod-394707")
DATASET_ID       = os.environ.get("DATASET_ID",       "cognito_prod_datamart")
CREDENTIALS_PATH = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "")
BQ_LOCATION      = os.environ.get("BQ_LOCATION",      "asia-south1")

# Table names — must match the constants in patent_thicket_analysis.py
SCORE_TABLE         = "Patent_Thicket_Score_Table"
CIRCUMVENTION_TABLE = "Circumvention_Table"

# ── Gemini Config ─────────────────────────────────────────────────────────────
GEMINI_MODEL = "gemini-2.5-flash"
API_KEY      = os.environ.get("GEMINI_API_KEY", "")

# ── Colour maps ───────────────────────────────────────────────────────────────
SCORE_COLOURS = {
    5: "C6EFCE", 4: "92D050", 3: "FFEB9C", 2: "FFC7CE", 1: "FF0000",
}
DIFFICULTY_COLOURS = {
    "Easy": "C6EFCE", "Moderate": "FFEB9C", "Difficult": "FFC7CE",
}
SCORE_LABELS = {
    5: "Exceptional - No meaningful secondary patent fence",
    4: "Strong - High probability of design-around",
    3: "Moderate - Requires structured strategy",
    2: "Weak - Circumvention and litigation become costly",
    1: "Poor - Dense Patent Thicket",
}

# ── LLM Narrative Prompt ──────────────────────────────────────────────────────
NARRATIVE_PROMPT = """
You are a pharmaceutical patent attorney and strategic consultant writing a concise
executive intelligence report for a 505(b)(2) drug development team.

Below is structured patent thicket analysis data for the drug "{drug}".
Write a professional 2-page narrative report section for this drug. Be precise,
analytical, and actionable. Do NOT use filler language. Keep each section concise
to fit within 2 printed pages.

--- STRUCTURED DATA ---
{data_json}
--- END DATA ---

Return ONLY a JSON object with these exact fields (no markdown, no preamble):
{{
  "executive_summary": "<2-3 sentences: overall IP landscape assessment and key takeaway>",
  "patent_landscape": "<1 short paragraph: describe the patent thicket - density, diversity, technology domains covered, what this means for a generic/505b2 entrant>",
  "thicket_score_narrative": "<2-3 sentences: interpret the final score {final_score}/5 ({score_label}) - what it means practically, key drivers including density score, diversity score, and validation %>",
  "circumvention_outlook": "<1 short paragraph: synthesise all design-around strategies across categories, highlight the most promising and most difficult categories, overall feasibility>",
  "key_risks": ["<risk 1>", "<risk 2>", "<risk 3>"],
  "conclusion": "<2-3 sentences: bottom-line assessment for the 505(b)(2) team>"
}}
"""

# ══════════════════════════════════════════════════════════════════════════════
# BigQuery helpers
# ══════════════════════════════════════════════════════════════════════════════

def _get_credentials():
    """Get credentials: use service account file if available, else default (Cloud Run)."""
    if CREDENTIALS_PATH and os.path.exists(CREDENTIALS_PATH):
        from google.oauth2 import service_account as sa
        return sa.Credentials.from_service_account_file(CREDENTIALS_PATH)
    return None  # Use ADC (Application Default Credentials)

def _bq_client():
    """Authenticated BigQuery client using service-account key or ADC."""
    try:
        from google.cloud import bigquery
    except ImportError:
        print("ERROR: google-cloud-bigquery not installed.\n"
              "Run: pip install google-cloud-bigquery db-dtypes --break-system-packages")
        sys.exit(1)

    credentials = _get_credentials()
    return bigquery.Client(project=PROJECT_ID, credentials=credentials,
                           location=BQ_LOCATION)


def _query_table(client, table_name: str):
    """Return a pandas DataFrame for an entire BQ table."""
    ref   = f"`{PROJECT_ID}.{DATASET_ID}.{table_name}`"
    query = f"SELECT DISTINCT * FROM {ref}"
    print(f"  Querying {ref} ...")
    try:
        df = client.query(query, location=BQ_LOCATION).to_dataframe()
        print(f"  -> {len(df)} rows, columns: {list(df.columns)}")
        return df
    except Exception as exc:
        print(f"ERROR querying {ref}: {exc}")
        sys.exit(1)


def _str(val) -> str:
    s = str(val).strip()
    return "" if s.lower() in ("nan", "none", "<na>", "nat") else s


def _int(val) -> int:
    try:
        return int(float(val))
    except (TypeError, ValueError):
        return 0


def _float(val) -> float:
    try:
        return float(val)
    except (TypeError, ValueError):
        return 0.0


def _score_label(score: int) -> str:
    return SCORE_LABELS.get(score, "N/A")


# ══════════════════════════════════════════════════════════════════════════════
# BigQuery narrative write-back
# ══════════════════════════════════════════════════════════════════════════════

def _narrative_columns() -> list:
    return [
        "narrative_executive_summary",
        "narrative_patent_landscape",
        "narrative_thicket_score",
        "narrative_circumvention_outlook",
        "narrative_key_risks",
        "narrative_conclusion",
        "created_at",
        "updated_at",
    ]


def _ensure_narrative_columns(client) -> None:
    from google.cloud import bigquery

    table_ref = f"{PROJECT_ID}.{DATASET_ID}.{SCORE_TABLE}"
    table     = client.get_table(table_ref)
    existing  = {field.name for field in table.schema}

    for col in _narrative_columns():
        if col in existing:
            print(f"    Column '{col}' already exists — skipping.")
            continue
        print(f"    Adding column '{col}' to {SCORE_TABLE} ...")
        col_type = "TIMESTAMP" if col in ("created_at", "updated_at") else "STRING"
        client.query(
            f"ALTER TABLE `{table_ref}` ADD COLUMN {col} {col_type}"
        ).result()
        print(f"    Column '{col}' added.")


def write_narrative_to_bigquery(drug_name: str, narrative: dict) -> None:
    from google.cloud import bigquery

    client    = _bq_client()
    table_ref = f"{PROJECT_ID}.{DATASET_ID}.{SCORE_TABLE}"

    print(f"  Writing narrative to BigQuery for drug: {drug_name}")
    _ensure_narrative_columns(client)

    key_risks_str = json.dumps(narrative.get("key_risks", []))

    sql = f"""
        UPDATE `{table_ref}`
        SET
            narrative_executive_summary   = @exec_summary,
            narrative_patent_landscape    = @patent_landscape,
            narrative_thicket_score       = @thicket_score,
            narrative_circumvention_outlook = @circ_outlook,
            narrative_key_risks           = @key_risks,
            narrative_conclusion          = @conclusion,
            created_at                    = CASE
                                              WHEN created_at IS NULL THEN @now
                                              ELSE created_at
                                            END,
            updated_at                    = @now
        WHERE Drug_Name = @drug_name
    """
    now_ts = datetime.now(tz=timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
    job_config = bigquery.QueryJobConfig(
        query_parameters=[
            bigquery.ScalarQueryParameter(
                "exec_summary",    "STRING", narrative.get("executive_summary",       "")),
            bigquery.ScalarQueryParameter(
                "patent_landscape","STRING", narrative.get("patent_landscape",         "")),
            bigquery.ScalarQueryParameter(
                "thicket_score",   "STRING", narrative.get("thicket_score_narrative",  "")),
            bigquery.ScalarQueryParameter(
                "circ_outlook",    "STRING", narrative.get("circumvention_outlook",    "")),
            bigquery.ScalarQueryParameter(
                "key_risks",       "STRING", key_risks_str),
            bigquery.ScalarQueryParameter(
                "conclusion",      "STRING", narrative.get("conclusion",               "")),
            bigquery.ScalarQueryParameter(
                "now",             "TIMESTAMP", now_ts),
            bigquery.ScalarQueryParameter(
                "drug_name",       "STRING", drug_name),
        ]
    )

    try:
        client.query(sql, job_config=job_config).result()
        print(f"    ✅ Narrative written for drug: {drug_name}")
    except Exception as e:
        print(f"    [ERROR] Failed to write narrative for drug '{drug_name}': {e}")
        raise


# ══════════════════════════════════════════════════════════════════════════════
# BigQuery reader
# ══════════════════════════════════════════════════════════════════════════════

def read_bq_data() -> dict:
    client = _bq_client()

    # ── Score table ───────────────────────────────────────────────────────────
    score_df  = _query_table(client, SCORE_TABLE)
    score_data: dict = {}

    for _, row in score_df.iterrows():
        drug         = _str(row.get("Drug_Name",    ""))
        jurisdiction = _str(row.get("Jurisdiction", ""))
        if not drug:
            continue

        if jurisdiction.lower().startswith("final score"):
            avg = _float(row.get("Final_Score", 0))
            if drug in score_data:
                score_data[drug]["avg_final_score"] = avg
            continue

        try:
            final_score = int(round(_float(row.get("Final_Score", 0))))
            if final_score not in (1, 2, 3, 4, 5):
                continue
        except (TypeError, ValueError):
            continue

        val_pct_raw = _float(row.get("Validation_Pct", 0))
        val_pct_str = f"{val_pct_raw:.1f}%"

        jur_entry = {
            "jurisdiction":      jurisdiction,
            "combined_total":    _int(row.get("Combined_Total",         0)),
            "adjusted_count":    _float(row.get("Adjusted_Count",       0)),
            "active_areas":      _int(row.get("Active_Technology_Areas",0)),
            "active_categories": _str(row.get("Active_Categories",      "")),
            "density_label":     _str(row.get("Density_Interpretation", "")),
            "diversity_label":   _str(row.get("Diversity_Interpretation","")),
            "density_score":     _int(row.get("Density_Score",  0)),
            "diversity_score":   _int(row.get("Diversity_Score",0)),
            "base_score":        _int(row.get("Base_Score",     0)),
            "validation_pct":    val_pct_str,
            "final_score":       final_score,
            "score_label":       _str(row.get("Score_Label",    "")),
        }

        if drug not in score_data:
            score_data[drug] = {"jurisdictions": [], "avg_final_score": None}
        score_data[drug]["jurisdictions"].append(jur_entry)

    for drug, sd in score_data.items():
        jurs = sd["jurisdictions"]
        if not jurs:
            continue

        if sd["avg_final_score"] is None:
            scores = [j["final_score"] for j in jurs]
            sd["avg_final_score"] = round(sum(scores) / len(scores), 1) if scores else 0

        avg_int = max(1, min(5, round(sd["avg_final_score"])))
        sd["final_score"]        = avg_int
        sd["score_label"]        = _score_label(avg_int)
        sd["total_combined"]     = sum(j["combined_total"] for j in jurs)
        sd["total_adjusted"]     = sum(j["adjusted_count"] for j in jurs)
        sd["total_active_areas"] = max(j["active_areas"]   for j in jurs)

        all_cats: set = set()
        for j in jurs:
            for c in j.get("active_categories", "").split(","):
                c = c.strip()
                if c:
                    all_cats.add(c)
        sd["active_categories"] = ", ".join(sorted(all_cats)) or "None"

        sd["density_label"]  = next(
            (j["density_label"]  for j in jurs if j["density_label"]),  "N/A")
        sd["diversity_label"] = next(
            (j["diversity_label"] for j in jurs if j["diversity_label"]), "N/A")

    # ── Circumvention table ───────────────────────────────────────────────────
    circ_df   = _query_table(client, CIRCUMVENTION_TABLE)
    circ_data: dict = {}

    for _, row in circ_df.iterrows():
        drug = _str(row.get("Drug_Name",       ""))
        cat  = _str(row.get("Patent_Category", ""))
        if not drug or not cat:
            continue

        if drug not in circ_data:
            circ_data[drug] = {}

        if cat not in circ_data[drug]:
            circ_data[drug][cat] = {
                "patents":                 _str(row.get("Patents",                    "")),
                "patent_count":            _int(row.get("Num_Patents",                0)),
                "overall_difficulty":      _str(row.get("Overall_Difficulty",         "")),
                "key_limitations":         _str(row.get("Key_Claim_Limitations",      "")),
                "white_space":             _str(row.get("White_Space_Opportunities",  "")),
                "fda_precedents":          _str(row.get("FDA_Precedents",             "")),
                "orange_book_gaps":        _str(row.get("Orange_Book_Gaps",           "")),
                "literature_alternatives": _str(row.get("Literature_Alternatives",    "")),
                "regulatory_viability":    _str(row.get("Regulatory_Viability",       "")),
                "summary":                 _str(row.get("Summary",                    "")),
                "strategies": [],
            }

        strategy = _str(row.get("Strategy", ""))
        if strategy and strategy.lower() != "no strategies identified":
            circ_data[drug][cat]["strategies"].append({
                "strategy":           strategy,
                "rationale":          _str(row.get("Rationale",          "")),
                "feasibility":        _str(row.get("Feasibility",        "")),
                "regulatory_pathway": _str(row.get("Regulatory_Pathway", "")),
                "prior_art_support":  _str(row.get("Prior_Art_Support",  "")),
            })

    # ── Merge ─────────────────────────────────────────────────────────────────
    all_drugs = set(score_data) | set(circ_data)
    merged: dict = {}
    for drug in all_drugs:
        merged[drug] = {
            "drug_name":          drug,
            "score_data":         score_data.get(drug, {}),
            "circumvention_data": circ_data.get(drug, {}),
        }

    return merged


# ══════════════════════════════════════════════════════════════════════════════
# LLM Narrative
# ══════════════════════════════════════════════════════════════════════════════

def call_gemini_for_narrative(drug_data: dict, api_key: str) -> dict:
    if not GEMINI_AVAILABLE:
        return _fallback_narrative(drug_data)

    client = genai.Client(api_key=api_key)

    sd          = drug_data.get("score_data", {})
    final_score = sd.get("final_score", "N/A")
    score_label = sd.get("score_label", "N/A")

    compact_circ = {}
    for cat, cd in drug_data.get("circumvention_data", {}).items():
        compact_circ[cat] = {
            "patent_count":        cd.get("patent_count", 0),
            "overall_difficulty":  cd.get("overall_difficulty", ""),
            "summary":             cd.get("summary", ""),
            "strategies": [
                {k: v for k, v in s.items()
                 if k in ("strategy", "feasibility", "rationale")}
                for s in cd.get("strategies", [])[:3]
            ],
            "fda_precedents":       cd.get("fda_precedents", ""),
            "regulatory_viability": cd.get("regulatory_viability", ""),
            "white_space":          cd.get("white_space", ""),
        }

    jurs = sd.get("jurisdictions", [])
    payload = {
        "drug_name": drug_data["drug_name"],
        "score_summary": {
            "final_score":       final_score,
            "score_label":       score_label,
            "avg_final_score":   sd.get("avg_final_score", "N/A"),
            "total_combined":    sd.get("total_combined",  0),
            "total_adjusted":    sd.get("total_adjusted",  0),
            "active_categories": sd.get("active_categories", ""),
            "density_label":     sd.get("density_label",  ""),
            "diversity_label":   sd.get("diversity_label",""),
            "jurisdictions": [
                {
                    "jurisdiction":    j["jurisdiction"],
                    "final_score":     j["final_score"],
                    "density_score":   j["density_score"],
                    "diversity_score": j["diversity_score"],
                    "base_score":      j["base_score"],
                    "validation_pct":  j["validation_pct"],
                    "combined_total":  j["combined_total"],
                    "adjusted_count":  j["adjusted_count"],
                    "active_areas":    j["active_areas"],
                }
                for j in jurs
            ],
        },
        "circumvention_by_category": compact_circ,
    }

    prompt = NARRATIVE_PROMPT.format(
        drug=drug_data["drug_name"],
        data_json=json.dumps(payload, indent=2),
        final_score=final_score,
        score_label=score_label,
    )

    try:
        resp = client.models.generate_content(
            model=GEMINI_MODEL, contents=prompt,
        )
        text = (resp.text or "").strip()
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$",          "", text).strip()
        return json.loads(text)
    except Exception as exc:
        print(f"  [WARN] Gemini call failed for {drug_data['drug_name']}: {exc}")
        return _fallback_narrative(drug_data)


def _fallback_narrative(drug_data: dict) -> dict:
    sd    = drug_data.get("score_data", {})
    drug  = drug_data["drug_name"]
    score = sd.get("final_score", "N/A")
    label = sd.get("score_label", "N/A")
    total = sd.get("total_combined", 0)
    jur_n = len(sd.get("jurisdictions", []))

    return {
        "executive_summary": (
            f"{drug} has {total} combined patents across {jur_n} jurisdiction(s). "
            f"Overall final score: {score}/5 ({label})."
        ),
        "patent_landscape": (
            f"The landscape for {drug} spans {total} combined patents across "
            f"{sd.get('total_active_areas', 0)} technology domain(s): "
            f"{sd.get('active_categories', 'N/A')}. "
            f"Density: {sd.get('density_label', 'N/A')}. "
            f"Diversity: {sd.get('diversity_label', 'N/A')}."
        ),
        "thicket_score_narrative": (
            f"The final score of {score}/5 ({label}) is the average across "
            f"{jur_n} jurisdiction(s), reflecting the relative ease of circumventing "
            f"the secondary patent portfolio surrounding {drug}."
        ),
        "circumvention_outlook": (
            "See Circumvention_Table in BigQuery for detailed design-around "
            "strategies per category."
        ),
        "key_risks": [
            "Review detailed circumvention analysis for category-specific risks.",
            "Monitor Orange Book listings for expiry and new filings.",
            "Assess freedom-to-operate before IND submission.",
        ],
        "conclusion": (
            f"The 505(b)(2) team should proceed with a structured design-around "
            f"analysis for {drug} given a final score of {score}/5."
        ),
    }


# ══════════════════════════════════════════════════════════════════════════════
# Word Document Builder
# ══════════════════════════════════════════════════════════════════════════════

def _hex_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _shd(cell, hex_color: str):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))

def _mar(cell, top=60, bottom=60, left=100, right=100):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'))

def _valign(cell, v="center"):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{v}"/>'))

def _cw(cell, dxa: int):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>'))

def _bdr(cell, color="CCCCCC", sz="4"):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top    w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left   w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right  w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'))

def _hdr_cell(cells, idx, text, dxa, bg="2F5496"):
    c = cells[idx]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shd(c, bg); _mar(c); _bdr(c, color=bg); _valign(c, "center"); _cw(c, dxa)

def _data_cell(cells, idx, text, dxa, bold=False, bg=None,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    c = cells[idx]; c.text = ""
    p = c.paragraphs[0]; p.alignment = align
    r = p.add_run(str(text) if text is not None else "")
    r.bold = bold; r.font.size = Pt(8.5); r.font.name = "Arial"
    if bg: _shd(c, bg)
    _mar(c); _bdr(c); _valign(c, "top"); _cw(c, dxa)

def _badge_cell(cells, idx, text, dxa, bg):
    c = cells[idx]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = True; r.font.size = Pt(8.5); r.font.name = "Arial"
    _shd(c, bg); _mar(c); _bdr(c, color=bg); _valign(c, "center"); _cw(c, dxa)

def _heading(doc, text, color="2F5496"):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.name = "Arial"; run.font.size = Pt(11)
        run.font.color.rgb = _hex_rgb(color)
    return p

def _body(doc, text, justify=False):
    p = doc.add_paragraph()
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(str(text or ""))
    r.font.size = Pt(9); r.font.name = "Arial"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(12)
    return p

def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.clear()
    r = p.add_run(str(text or ""))
    r.font.size = Pt(9); r.font.name = "Arial"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(12)
    return p

def _divider(doc, color="2F5496"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'))

def _banner(doc, drug_name):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="1F3864" w:val="clear"/>'))
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="2F5496"/>'
        f'</w:pBdr>'))
    r = p.add_run(drug_name)
    r.bold = True; r.font.size = Pt(16); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def _subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(8); r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88); r.italic = True

def _add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Patent Thicket Analysis - Confidential  |  Page ")
    r1.font.size = Pt(7); r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r3 = fp.add_run()
    r3._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r3.font.size = Pt(7); r3.font.name = "Arial"
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fp.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


# ══════════════════════════════════════════════════════════════════════════════
# PDF Conversion — pip-only, no LibreOffice required
# Uses: mammoth (docx → HTML) + xhtml2pdf (HTML → PDF)
# Install: pip install mammoth xhtml2pdf --break-system-packages
# ══════════════════════════════════════════════════════════════════════════════

# Hex colour → CSS colour helper
def _hex_to_css(h: str) -> str:
    h = h.lstrip("#")
    return f"#{h}" if len(h) == 6 else "#000000"

# Map score int → CSS background colour
def _score_css_bg(score) -> str:
    return _hex_to_css(SCORE_COLOURS.get(score, "DDDDDD"))

def _difficulty_css_bg(difficulty: str) -> str:
    return _hex_to_css(DIFFICULTY_COLOURS.get(difficulty, "FFFFFF"))


def _build_html_for_drug(entry: dict) -> str:
    """
    Build a self-contained HTML page for one drug entry.
    xhtml2pdf renders this to PDF with full table + colour support.
    """
    drug_name = entry["drug_name"]
    sd        = entry.get("score_data", {})
    circ      = entry.get("circumvention_data", {})
    narr      = entry.get("narrative", {})

    score       = sd.get("final_score", "N/A")
    score_bg    = _score_css_bg(score) if isinstance(score, int) else "#DDDDDD"
    avg_final   = sd.get("avg_final_score", "N/A")
    jurs        = sd.get("jurisdictions", [])
    date_str    = datetime.now().strftime("%B %d, %Y")

    # ── Summary KPI table ─────────────────────────────────────────────────────
    kpi_rows = f"""
    <tr>
      <td class="badge" style="background:{score_bg}">{score} / 5</td>
      <td class="center">{avg_final}</td>
      <td class="center">{sd.get("total_combined","N/A")}</td>
      <td class="center">{len(jurs)}</td>
      <td class="center">{sd.get("density_label","N/A")}</td>
    </tr>
    """

    # ── Jurisdiction breakdown ────────────────────────────────────────────────
    jur_rows = ""
    for j in jurs:
        jf     = j.get("final_score", 0)
        jf_bg  = _score_css_bg(jf)
        jur_rows += f"""
        <tr>
          <td><b>{j["jurisdiction"]}</b></td>
          <td class="center">{j.get("density_score","")}</td>
          <td class="center">{j.get("diversity_score","")}</td>
          <td class="center">{j.get("base_score","")}</td>
          <td class="center">{j.get("validation_pct","")}</td>
          <td class="badge" style="background:{jf_bg}">{jf}</td>
        </tr>"""

    jur_section = ""
    if jurs:
        jur_section = f"""
        <h2>Score Breakdown by Jurisdiction</h2>
        <table>
          <thead><tr>
            <th>Jurisdiction</th><th>Density</th><th>Diversity</th>
            <th>Base</th><th>Validation</th><th>Final</th>
          </tr></thead>
          <tbody>{jur_rows}</tbody>
        </table>"""

    # ── Design-around summary table ───────────────────────────────────────────
    cat_rows = ""
    for cat, cd in circ.items():
        diff    = cd.get("overall_difficulty", "N/A")
        diff_bg = _difficulty_css_bg(diff)
        sm      = cd.get("summary", "") or "No summary available"
        if sm.lower() in ("nan", "none", ""):
            sm = "No summary available"
        cat_rows += f"""
        <tr>
          <td>{cat}</td>
          <td class="center"><b>{cd.get("patent_count", 0)}</b></td>
          <td class="badge" style="background:{diff_bg}">{diff}</td>
          <td>{sm}</td>
        </tr>"""

    cat_section = ""
    if circ:
        cat_section = f"""
        <h2>Design-Around Summary by Category</h2>
        <table>
          <thead><tr>
            <th style="width:22%">Category</th>
            <th style="width:9%">Patents</th>
            <th style="width:15%">Difficulty</th>
            <th>Summary</th>
          </tr></thead>
          <tbody>{cat_rows}</tbody>
        </table>"""

    # ── Key risks ─────────────────────────────────────────────────────────────
    risks     = narr.get("key_risks", [])
    risk_html = "".join(f"<li>{r}</li>" for r in risks)
    risk_section = f"<h2>Key Risks</h2><ul>{risk_html}</ul>" if risks else ""

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{
    size: A4;
    margin: 1.8cm 2cm 2.4cm 2cm;
  }}
  .pdf-footer {{
    position: fixed;
    bottom: 0;
    left: 0;
    right: 0;
    text-align: center;
    font-size: 7pt;
    color: #888888;
    font-family: Arial, sans-serif;
    border-top: 1px solid #CCCCCC;
    padding-top: 3px;
  }}
  body {{
    font-family: Arial, sans-serif;
    font-size: 9pt;
    color: #222222;
    margin: 0;
    padding: 0;
  }}
  .banner {{
    background: #1F3864;
    color: #ffffff;
    font-size: 16pt;
    font-weight: bold;
    padding: 10px 14px 8px 14px;
    border-bottom: 3px solid #2F5496;
    margin-bottom: 2px;
  }}
  .subtitle {{
    font-size: 8pt;
    color: #888888;
    font-style: italic;
    margin-bottom: 8px;
    padding-left: 2px;
  }}
  h2 {{
    color: #2F5496;
    font-size: 11pt;
    font-family: Arial, sans-serif;
    margin: 10px 0 3px 0;
    padding: 0;
    border-bottom: 1px solid #2F5496;
    padding-bottom: 2px;
  }}
  p {{
    margin: 0 0 5px 0;
    line-height: 1.35;
    text-align: justify;
  }}
  ul {{
    margin: 2px 0 6px 0;
    padding-left: 18px;
  }}
  li {{
    margin-bottom: 2px;
    line-height: 1.3;
  }}
  table {{
    border-collapse: collapse;
    width: 100%;
    margin: 5px 0 8px 0;
    font-size: 8.5pt;
  }}
  thead tr th {{
    background: #2F5496;
    color: #ffffff;
    font-weight: bold;
    text-align: center;
    padding: 4px 6px;
    border: 1px solid #2F5496;
  }}
  tbody tr td {{
    border: 1px solid #CCCCCC;
    padding: 4px 6px;
    vertical-align: top;
  }}
  .center {{ text-align: center; }}
  .badge {{
    text-align: center;
    font-weight: bold;
    border: none !important;
  }}
  .divider {{
    border: none;
    border-bottom: 2px solid #2F5496;
    margin: 10px 0 6px 0;
  }}
  .page-break {{ page-break-before: always; }}
</style>
</head>
<body>

<div class="banner">{drug_name}</div>
<div class="subtitle">Patent Thicket &amp; Circumvention Feasibility Analysis &nbsp;|&nbsp; {date_str}</div>

<!-- KPI Summary Table -->
<table>
  <thead><tr>
    <th>Final Score</th><th>Avg Score</th>
    <th>Combined Patents</th><th>Jurisdictions</th><th>Density</th>
  </tr></thead>
  <tbody>{kpi_rows}</tbody>
</table>

{jur_section}

<h2>Executive Summary</h2>
<p>{narr.get("executive_summary","")}</p>

<h2>Patent Landscape</h2>
<p>{narr.get("patent_landscape","")}</p>

<h2>Final Score: {score}/5 — {sd.get("score_label","")}</h2>
<p>{narr.get("thicket_score_narrative","")}</p>

<h2>Circumvention Outlook</h2>
<p>{narr.get("circumvention_outlook","")}</p>

{cat_section}

{risk_section}

<h2>Conclusion</h2>
<p>{narr.get("conclusion","")}</p>

<hr class="divider">

<div class="pdf-footer">Patent Thicket Analysis - Confidential</div>

</body>
</html>"""
    return html


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert a .docx file to PDF using:
      mammoth   — converts .docx → HTML  (preserves structure)
      xhtml2pdf — renders HTML → PDF     (pure pip, works on Windows/Linux/macOS)

    Install:
        pip install mammoth xhtml2pdf --break-system-packages
    """
    try:
        import mammoth       # noqa: F401
        from xhtml2pdf import pisa  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "mammoth and/or xhtml2pdf are not installed.\n"
            "Run: pip install mammoth xhtml2pdf --break-system-packages"
        )

    pdf_path  = os.path.splitext(docx_path)[0] + ".pdf"

    # Step 1: docx → HTML via mammoth
    import mammoth
    style_map = """
        p[style-name='Heading 2'] => h2:fresh
        r[style-name='Strong']    => strong
    """
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f, style_map=style_map)
    raw_html = result.value

    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{ margin: 1.8cm 2cm 2cm 2cm; }}
  body  {{ font-family: Arial, sans-serif; font-size: 9pt; color: #222; margin:0; padding:0; }}
  h1    {{ color: #1F3864; font-size: 14pt; border-bottom: 3px solid #2F5496; padding-bottom:4px; }}
  h2    {{ color: #2F5496; font-size: 11pt; border-bottom: 1px solid #2F5496; padding-bottom:2px; margin:10px 0 3px 0; }}
  p     {{ margin: 0 0 5px 0; line-height: 1.35; text-align: justify; }}
  ul    {{ margin: 2px 0 6px 0; padding-left: 18px; }}
  li    {{ margin-bottom: 2px; line-height: 1.3; }}
  table {{ border-collapse: collapse; width: 100%; margin: 5px 0 8px 0; font-size: 8.5pt; }}
  th    {{ background: #2F5496; color: #ffffff; font-weight: bold; text-align: center; padding: 4px 6px; border: 1px solid #2F5496; }}
  td    {{ border: 1px solid #CCCCCC; padding: 4px 6px; vertical-align: top; }}
  hr    {{ border: none; border-bottom: 2px solid #2F5496; margin: 10px 0 6px 0; }}
</style>
</head>
<body>
{raw_html}
</body>
</html>"""

    # Step 2: HTML → PDF via xhtml2pdf
    from xhtml2pdf import pisa
    with open(pdf_path, "wb") as pdf_file:
        result = pisa.CreatePDF(full_html.encode("utf-8"), dest=pdf_file,
                                encoding="utf-8")
    if result.err:
        raise RuntimeError(f"xhtml2pdf conversion error code: {result.err}")

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF not created at expected path: {pdf_path}")
    return pdf_path


def convert_entry_to_pdf_direct(entry: dict, pdf_path: str) -> str:
    """
    Bypass docx entirely — render drug entry directly from Python data
    → styled HTML → PDF via xhtml2pdf.
    Pure pip install, works on Windows, Linux, and macOS with no system deps.

    Install: pip install xhtml2pdf --break-system-packages
    """
    try:
        from xhtml2pdf import pisa  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "xhtml2pdf is not installed.\n"
            "Run: pip install xhtml2pdf --break-system-packages"
        )

    html = _build_html_for_drug(entry)

    from xhtml2pdf import pisa
    with open(pdf_path, "wb") as pdf_file:
        result = pisa.CreatePDF(html.encode("utf-8"), dest=pdf_file,
                                encoding="utf-8")
    if result.err:
        raise RuntimeError(f"xhtml2pdf conversion error code: {result.err}")

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF not created at expected path: {pdf_path}")
    return pdf_path


# ══════════════════════════════════════════════════════════════════════════════
# GCS Upload
# ══════════════════════════════════════════════════════════════════════════════

def upload_to_gcs(local_pdf: str, drug_name: str,
                  bucket_name: str = "cognito-gcs") -> str:
    try:
        from google.cloud import storage
    except ImportError:
        raise RuntimeError(
            "google-cloud-storage not installed.\n"
            "Run: pip install google-cloud-storage --break-system-packages"
        )

    safe_drug = drug_name.replace("/", "_").replace(" ", "_")
    blob_name = f"Cognito_new/reports/{safe_drug}/Patent_Thicket_Analysis.pdf"

    credentials = _get_credentials()
    client = storage.Client(project=PROJECT_ID, credentials=credentials)

    bucket = client.bucket(bucket_name)
    blob   = bucket.blob(blob_name)
    blob.upload_from_filename(local_pdf, content_type="application/pdf")

    gcs_uri = f"gs://{bucket_name}/{blob_name}"
    print(f"  Uploaded → {gcs_uri}")
    return gcs_uri


# ══════════════════════════════════════════════════════════════════════════════
# Word Document Builder (unchanged from original)
# ══════════════════════════════════════════════════════════════════════════════

def build_document(drugs_list: list, output_path: str):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Emu(12240 * 635)
    sec.page_height   = Emu(15840 * 635)
    sec.top_margin    = Emu(720   * 635)
    sec.bottom_margin = Emu(720   * 635)
    sec.left_margin   = Emu(900   * 635)
    sec.right_margin  = Emu(900   * 635)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    _add_footer(sec)

    for i, entry in enumerate(drugs_list):
        if i > 0:
            doc.add_page_break()

        drug_name = entry["drug_name"]
        sd        = entry.get("score_data", {})
        circ      = entry.get("circumvention_data", {})
        narr      = entry.get("narrative", {})

        score       = sd.get("final_score", "N/A")
        score_color = (SCORE_COLOURS.get(score, "DDDDDD")
                       if isinstance(score, int) else "DDDDDD")
        avg_final   = sd.get("avg_final_score", "N/A")
        jurs        = sd.get("jurisdictions", [])

        _banner(doc, drug_name)
        _subtitle(doc,
            "Patent Thicket & Circumvention Feasibility Analysis  |  "
            + datetime.now().strftime("%B %d, %Y"))

        tbl = doc.add_table(rows=2, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        _hdr_cell(hdr, 0, "Final Score",      1870)
        _hdr_cell(hdr, 1, "Avg Score",        1870)
        _hdr_cell(hdr, 2, "Combined Patents", 1870)
        _hdr_cell(hdr, 3, "Jurisdictions",    1870)
        _hdr_cell(hdr, 4, "Density",          1870)
        dat = tbl.rows[1].cells
        _badge_cell(dat, 0, f"{score} / 5",                 1870, score_color)
        _data_cell (dat, 1, avg_final,                      1870, bold=True, bg="F2F2F2",
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _data_cell (dat, 2, sd.get("total_combined","N/A"), 1870, bold=True, bg="F2F2F2",
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _data_cell (dat, 3, len(jurs),                      1870, bold=True, bg="F2F2F2",
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _data_cell (dat, 4, sd.get("density_label","N/A"),  1870, bg="F2F2F2",
                    align=WD_ALIGN_PARAGRAPH.CENTER)

        if jurs:
            _heading(doc, "Score Breakdown by Jurisdiction")
            jtbl = doc.add_table(rows=1 + len(jurs), cols=6)
            jtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            jh = jtbl.rows[0].cells
            _hdr_cell(jh, 0, "Jurisdiction", 1500)
            _hdr_cell(jh, 1, "Density",      1300)
            _hdr_cell(jh, 2, "Diversity",    1300)
            _hdr_cell(jh, 3, "Base",         1200)
            _hdr_cell(jh, 4, "Validation",   1500)
            _hdr_cell(jh, 5, "Final",        1200)
            for ji, j in enumerate(jurs, start=1):
                jc = jtbl.rows[ji].cells
                jf = j.get("final_score", 0)
                _data_cell (jc, 0, j["jurisdiction"],           1500, bold=True)
                _data_cell (jc, 1, j.get("density_score",  ""), 1300,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
                _data_cell (jc, 2, j.get("diversity_score",""), 1300,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
                _data_cell (jc, 3, j.get("base_score",     ""), 1200,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
                _data_cell (jc, 4, j.get("validation_pct", ""), 1500,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
                _badge_cell(jc, 5, jf, 1200, SCORE_COLOURS.get(jf, "DDDDDD"))

        _heading(doc, "Executive Summary")
        _body(doc, narr.get("executive_summary", ""), justify=True)

        _heading(doc, "Patent Landscape")
        _body(doc, narr.get("patent_landscape", ""), justify=True)

        _heading(doc, f"Final Score: {score}/5 - {sd.get('score_label','')}")
        _body(doc, narr.get("thicket_score_narrative", ""), justify=True)

        _heading(doc, "Circumvention Outlook")
        _body(doc, narr.get("circumvention_outlook", ""), justify=True)

        cats = list(circ.items())
        if cats:
            _heading(doc, "Design-Around Summary by Category")
            ctbl = doc.add_table(rows=1 + len(cats), cols=4)
            ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            ch = ctbl.rows[0].cells
            _hdr_cell(ch, 0, "Category",   2500)
            _hdr_cell(ch, 1, "Patents",    1000)
            _hdr_cell(ch, 2, "Difficulty", 1700)
            _hdr_cell(ch, 3, "Summary",    4160)
            for ri, (cat, cd) in enumerate(cats, start=1):
                rc   = ctbl.rows[ri].cells
                diff = cd.get("overall_difficulty", "N/A")
                sm   = cd.get("summary", "") or "No summary available"
                if sm.lower() in ("nan", "none", ""):
                    sm = "No summary available"
                _data_cell (rc, 0, cat,                        2500)
                _data_cell (rc, 1, cd.get("patent_count", 0),  1000, bold=True,
                            bg="F2F2F2", align=WD_ALIGN_PARAGRAPH.CENTER)
                _badge_cell(rc, 2, diff, 1700, DIFFICULTY_COLOURS.get(diff, "FFFFFF"))
                _data_cell (rc, 3, sm,   4160)

        risks = narr.get("key_risks", [])
        if risks:
            _heading(doc, "Key Risks")
            for r in risks:
                _bullet(doc, r)

        _heading(doc, "Conclusion")
        _body(doc, narr.get("conclusion", ""), justify=True)
        _divider(doc)

    doc.save(output_path)
    print(f"  Saved: {output_path}")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

# Set to True  → render PDF directly from Python data via xhtml2pdf (works on Windows, recommended)
# Set to False → go docx → mammoth HTML → xhtml2pdf PDF (closer to original docx flow)
USE_DIRECT_HTML_RENDER = True


def generate_report(api_key: str = "", gcs_bucket: str = "cognito-gcs"):
    if not api_key:
        api_key = API_KEY

    if not api_key:
        raise EnvironmentError(
            "GEMINI_API_KEY is not set. Add it to your .env file:\n"
            "  GEMINI_API_KEY=your-key"
        )
    # Credentials resolved automatically (service account file or ADC)

    print(f"\n{'='*60}")
    print(f"Patent Thicket Report Generator - BigQuery Edition")
    print(f"Project    : {PROJECT_ID}")
    print(f"Dataset    : {DATASET_ID}")
    print(f"Location   : {BQ_LOCATION}")
    print(f"GCS bucket : gs://{gcs_bucket}/Cognito_new/reports/{{drug_name}}/Patent_Thicket_Analysis.pdf")
    print(f"PDF engine : {'xhtml2pdf (direct HTML)' if USE_DIRECT_HTML_RENDER else 'mammoth + xhtml2pdf'}")
    print(f"{'='*60}")

    print("\n[1/4] Reading from BigQuery...")
    drugs_data = read_bq_data()
    if not drugs_data:
        print("ERROR: No drug data found in BigQuery tables.")
        sys.exit(1)
    print(f"  Found {len(drugs_data)} drug(s): {', '.join(sorted(drugs_data))}")

    print("\n[2/4] Generating narratives and writing back to BigQuery...")

    drugs_list = []
    for drug_name, drug_data in sorted(drugs_data.items()):
        print(f"\n  {drug_name} ...")
        if api_key and GEMINI_AVAILABLE:
            narrative = call_gemini_for_narrative(drug_data, api_key)
        else:
            narrative = _fallback_narrative(drug_data)

        try:
            write_narrative_to_bigquery(drug_name, narrative)
        except Exception as exc:
            print(f"  [WARN] BigQuery write-back failed for '{drug_name}': {exc}")

        drugs_list.append({
            "drug_name":          drug_name,
            "score_data":         drug_data.get("score_data", {}),
            "circumvention_data": drug_data.get("circumvention_data", {}),
            "narrative":          narrative,
        })

    print(f"\n[3/4] Building PDF reports...")
    uploaded_uris = []
    with tempfile.TemporaryDirectory() as tmpdir:
        for entry in drugs_list:
            drug_name = entry["drug_name"]
            safe_drug = drug_name.replace("/", "_").replace(" ", "_")
            pdf_path  = os.path.join(tmpdir, f"{safe_drug}.pdf")

            print(f"\n  [{drug_name}]")

            if USE_DIRECT_HTML_RENDER:
                # Fast path: data → HTML → PDF (no docx intermediate)
                print(f"    Rendering PDF (xhtml2pdf direct) ...", end=" ", flush=True)
                try:
                    convert_entry_to_pdf_direct(entry, pdf_path)
                    print("done")
                except RuntimeError as exc:
                    print(f"FAILED\n    ERROR: {exc}")
                    continue
            else:
                # Docx path: data → docx → mammoth HTML → PDF
                docx_path = os.path.join(tmpdir, f"{safe_drug}.docx")
                print(f"    Building DOCX ...", end=" ", flush=True)
                build_document([entry], docx_path)
                print("done")

                print(f"    Converting to PDF (mammoth + xhtml2pdf) ...", end=" ", flush=True)
                try:
                    pdf_path = convert_docx_to_pdf(docx_path)
                    print("done")
                except RuntimeError as exc:
                    print(f"FAILED\n    ERROR: {exc}")
                    continue

            print(f"\n[4/4] Uploading to GCS ...")
            print(f"    Uploading to GCS ...", end=" ", flush=True)
            try:
                uri = upload_to_gcs(pdf_path, drug_name, bucket_name=gcs_bucket)
                uploaded_uris.append(uri)
                print("done")
            except Exception as exc:
                print(f"FAILED\n    ERROR: {exc}")

    print(f"\n  {len(uploaded_uris)} PDF(s) uploaded successfully.")
    for uri in uploaded_uris:
        print(f"    {uri}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Patent Thicket Report Generator - BigQuery Edition"
    )
    parser.add_argument("--api-key",  default="",
                        help="Gemini API key (overrides GEMINI_API_KEY env var)")
    parser.add_argument("--project",  default="", help="GCP Project ID")
    parser.add_argument("--dataset",  default="", help="BigQuery Dataset ID")
    parser.add_argument("--key-file", default="",
                        help="Service account JSON path (overrides GOOGLE_APPLICATION_CREDENTIALS)")
    parser.add_argument("--location", default="", help="BigQuery location")
    parser.add_argument("--bucket",   default="cognito-gcs",
                        help="GCS bucket name (default: cognito-gcs)")
    parser.add_argument("--no-direct-render", action="store_true",
                        help="Use docx→mammoth→xhtml2pdf instead of direct HTML render")
    args = parser.parse_args()

    if args.project:          PROJECT_ID            = args.project
    if args.dataset:          DATASET_ID            = args.dataset
    if args.key_file:         CREDENTIALS_PATH      = args.key_file
    if args.location:         BQ_LOCATION           = args.location
    if args.no_direct_render: USE_DIRECT_HTML_RENDER = False

    generate_report(args.api_key, gcs_bucket=args.bucket)
