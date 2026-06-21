"""
generate_report.py
──────────────────
Reads scored data from BigQuery tables produced by the Patent Legal Robustness Scorer
and generates a professional detailed Word (.docx) report using Gemini 2.5 Flash
for narrative generation and python-docx for document creation.

Input tables (BigQuery):
  - patent_strength_table
  - patent_strength_country_score_table

Usage:
    pip install python-docx pandas google-cloud-bigquery google-genai db-dtypes google-cloud-storage python-dotenv --break-system-packages
    # Create a .env file with:
    #   GEMINI_API_KEY=your-key
    #   GCS_CREDENTIALS=/path/to/service-account.json
    python generate_report.py --output report.docx
"""

import os
import re
import json
import argparse
from datetime import datetime, timezone
from pathlib import Path

# ── Load .env ─────────────────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # Not needed on Cloud Run

import pandas as pd
from google.cloud import bigquery
from google.oauth2 import service_account
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from google import genai as genai_client
from google.genai import types

# ── Config ────────────────────────────────────────────────────────────────────
# Loaded from environment / .env file
API_KEY        = os.environ.get("GEMINI_API_KEY", "")
CREDENTIALS_PATH = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "")

MODEL = "gemini-2.5-flash"

BQ_PROJECT_ID = "cognito-prod-394707"
BQ_DATASET_ID = "cognito_prod_datamart"
BQ_LOCATION   = "asia-south1"

BQ_STRENGTH_TABLE      = "patent_strength_table"
BQ_COUNTRY_SCORE_TABLE = "patent_strength_country_score_table"

# ── GCS Configuration ─────────────────────────────────────────────────────────
GCS_BUCKET    = "cognito-gcs"
GCS_BASE_PATH = "Cognito_new/reports"
GCS_FILE_NAME = "Patent_Strength_Analysis.docx"

SCORE_COLOR_MAP = {
    1: RGBColor(0x00, 0x80, 0x00),  # Green
    2: RGBColor(0x4C, 0xAF, 0x50),  # Light green
    3: RGBColor(0xCC, 0x99, 0x00),  # Amber
    4: RGBColor(0xE6, 0x51, 0x00),  # Orange-red
    5: RGBColor(0xCC, 0x00, 0x00),  # Red
}

SCORE_LABEL = {
    1: "Very Robust",
    2: "Robust",
    3: "Moderate",
    4: "Vulnerable",
    5: "Highly Vulnerable",
}

REPORT_TITLE = "Patent Strength Scoring Analysis"


# ── Gemini helper ─────────────────────────────────────────────────────────────

def call_gemini(prompt: str) -> str:
    """Call Gemini 2.5 Flash and return plain text."""
    client = genai_client.Client(api_key=API_KEY)
    config = types.GenerateContentConfig(temperature=0.3)
    response = client.models.generate_content(
        model=MODEL, contents=prompt, config=config
    )
    return response.text.strip() if response.text else ""


def _extract_json(text: str):
    """Extract JSON from a Gemini response that may contain markdown fences."""
    text = re.sub(r"^```(?:json)?", "", text.strip()).strip()
    text = re.sub(r"```$", "", text).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass
    return None


# ── Data loading ──────────────────────────────────────────────────────────────

def _get_credentials():
    """Get credentials: use service account file if available, else default (Cloud Run)."""
    if CREDENTIALS_PATH and os.path.exists(CREDENTIALS_PATH):
        return service_account.Credentials.from_service_account_file(CREDENTIALS_PATH)
    return None  # Use ADC (Application Default Credentials)


def _bq_client():
    credentials = _get_credentials()
    return bigquery.Client(
        project=BQ_PROJECT_ID,
        credentials=credentials,
        location=BQ_LOCATION,
    )


def load_from_bigquery() -> dict:
    """Read scored data from BigQuery tables produced by the scorer."""
    client = _bq_client()
    data = {}

    # ── patent_strength_table ─────────────────────────────────────────────────
    print(f"Loading {BQ_STRENGTH_TABLE} from BigQuery...")
    try:
        q = f"SELECT DISTINCT * FROM `{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_STRENGTH_TABLE}`"
        df_final = client.query(q).to_dataframe()
        df_final = df_final.rename(columns={
            "drug_name":           "Drug Name",
            "patent_number":       "Patent Number",
            "patent_type":         "Patent Type",
            "jurisdiction":        "Jurisdiction",
            "country_weight":      "Country Weight",
            "core_inventive_step": "Core Inventive Step",
            "sf1_score":           "SF1 Score",
            "sf2_score":           "SF2 Score",
            "sf3_score":           "SF3 Score",
            "sf4_score":           "SF4 Score",
            "weighted_score":      "Weighted Final Score",
            "key_vulnerabilities": "Key Vulnerabilities",
            "key_strengths":       "Key Strengths",
            "filing_date":         "Filing Date",
            "grant_date":          "Grant Date",
            "sf1_label":           "SF1 Label",
            "sf1_score_reason":    "SF1 Score Reason",
            "sf1_reasoning":       "SF1 Reasoning",
            "sf1_key_finding_1":   "SF1 Key Finding 1",
            "sf1_key_finding_2":   "SF1 Key Finding 2",
            "sf1_key_finding_3":   "SF1 Key Finding 3",
            "sf1_chunks_used":     "SF1 Chunks Used",
            "sf2_label":           "SF2 Label",
            "sf2_score_reason":    "SF2 Score Reason",
            "sf2_reasoning":       "SF2 Reasoning",
            "sf2_key_finding_1":   "SF2 Key Finding 1",
            "sf2_key_finding_2":   "SF2 Key Finding 2",
            "sf2_key_finding_3":   "SF2 Key Finding 3",
            "sf2_chunks_used":     "SF2 Chunks Used",
            "sf3_label":           "SF3 Label",
            "sf3_score_reason":    "SF3 Score Reason",
            "sf3_reasoning":       "SF3 Reasoning",
            "sf3_key_finding_1":   "SF3 Key Finding 1",
            "sf3_key_finding_2":   "SF3 Key Finding 2",
            "sf3_key_finding_3":   "SF3 Key Finding 3",
            "sf3_chunks_used":     "SF3 Chunks Used",
            "sf4_label":           "SF4 Label",
            "sf4_score_reason":    "SF4 Score Reason",
            "sf4_reasoning":       "SF4 Reasoning",
            "sf4_key_finding_1":   "SF4 Key Finding 1",
            "sf4_key_finding_2":   "SF4 Key Finding 2",
            "sf4_key_finding_3":   "SF4 Key Finding 3",
            "sf4_chunks_used":     "SF4 Chunks Used",
            "created_at":          "Created At",
            "updated_at":          "Updated At",
        })
        data["final"] = df_final
        print(f"  Loaded {len(df_final)} rows.")
    except Exception as e:
        print(f"  WARNING: Could not load {BQ_STRENGTH_TABLE}: {e}")
        data["final"] = pd.DataFrame()

    # ── patent_strength_country_score_table ───────────────────────────────────
    print(f"Loading {BQ_COUNTRY_SCORE_TABLE} from BigQuery...")
    try:
        q = f"SELECT DISTINCT * FROM `{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_COUNTRY_SCORE_TABLE}`"
        df_country = client.query(q).to_dataframe()
        df_country = df_country.rename(columns={
            "drug_name":              "Drug Name",
            "jurisdiction":           "Jurisdiction",
            "country_name":           "Country Name",
            "country_weight":         "Country Weight",
            "patent_count":           "# Patents",
            "avg_weighted_score":     "Avg Weighted Score",
            "country_weighted_score": "Country Weighted Score",
            "final_patent_score":     "Final Patent Score (Drug Total)",
        })
        data["country_scores"] = df_country
        print(f"  Loaded {len(df_country)} rows.")
    except Exception as e:
        print(f"  WARNING: Could not load {BQ_COUNTRY_SCORE_TABLE}: {e}")
        data["country_scores"] = pd.DataFrame()

    return data


def compute_statistics(df: pd.DataFrame) -> dict:  # noqa: C901
    """Compute summary statistics from the Final Output sheet."""
    stats = {
        "total_patents": len(df),
        "drugs": [],
        "avg_weighted": None,
        "score_distribution": {},
        "most_vulnerable": [],
        "most_robust": [],
        "per_drug_stats": {},
        "highest_score_patent": None,
        "highest_score_per_jurisdiction": {},
    }
    if df.empty:
        return stats

    stats["drugs"] = [str(d) for d in df["Drug Name"].dropna().unique().tolist()]

    ws_col = "Weighted Final Score"
    if ws_col in df.columns:
        numeric = pd.to_numeric(df[ws_col], errors="coerce").dropna()
        if len(numeric):
            stats["avg_weighted"] = float(round(numeric.mean(), 2))
            for s in range(1, 6):
                stats["score_distribution"][s] = int((numeric.round() == s).sum())
            top_vuln = df.loc[numeric.nlargest(3).index]
            stats["most_vulnerable"] = top_vuln[["Drug Name", "Patent Number", ws_col]].to_dict("records")
            top_rob = df.loc[numeric.nsmallest(3).index]
            stats["most_robust"] = top_rob[["Drug Name", "Patent Number", ws_col]].to_dict("records")

            highest_idx = numeric.idxmax()
            highest_row = df.loc[highest_idx]
            stats["highest_score_patent"] = {
                "drug": str(highest_row.get("Drug Name", "N/A")),
                "patent_number": str(highest_row.get("Patent Number", "N/A")),
                "patent_type": str(highest_row.get("Patent Type", "N/A")),
                "jurisdiction": str(highest_row.get("Jurisdiction", "N/A")),
                "weighted_score": float(numeric[highest_idx]),
                "sf1": highest_row.get("SF1 Score", "N/A"),
                "sf2": highest_row.get("SF2 Score", "N/A"),
                "sf3": highest_row.get("SF3 Score", "N/A"),
                "sf4": highest_row.get("SF4 Score", "N/A"),
                "core_step": str(highest_row.get("Core Inventive Step", "N/A")),
                "vulnerabilities": str(highest_row.get("Key Vulnerabilities", "N/A")),
                "strengths": str(highest_row.get("Key Strengths", "N/A")),
            }

            if "Jurisdiction" in df.columns:
                per_jurisdiction = {}
                for jur in df["Jurisdiction"].dropna().unique():
                    jur_df = df[df["Jurisdiction"] == jur]
                    jur_numeric = pd.to_numeric(jur_df[ws_col], errors="coerce")
                    jur_valid = jur_numeric.dropna()
                    if len(jur_valid):
                        jur_highest_idx = jur_valid.idxmax()
                        jur_row = df.loc[jur_highest_idx]
                        per_jurisdiction[str(jur)] = {
                            "drug": str(jur_row.get("Drug Name", "N/A")),
                            "patent_number": str(jur_row.get("Patent Number", "N/A")),
                            "patent_type": str(jur_row.get("Patent Type", "N/A")),
                            "jurisdiction": str(jur),
                            "weighted_score": float(jur_valid[jur_highest_idx]),
                            "sf1": jur_row.get("SF1 Score", "N/A"),
                            "sf2": jur_row.get("SF2 Score", "N/A"),
                            "sf3": jur_row.get("SF3 Score", "N/A"),
                            "sf4": jur_row.get("SF4 Score", "N/A"),
                            "core_step": str(jur_row.get("Core Inventive Step", "N/A")),
                            "vulnerabilities": str(jur_row.get("Key Vulnerabilities", "N/A")),
                            "strengths": str(jur_row.get("Key Strengths", "N/A")),
                        }
                stats["highest_score_per_jurisdiction"] = per_jurisdiction

        for drug in stats["drugs"]:
            drug_df = df[df["Drug Name"] == drug]
            drug_numeric = pd.to_numeric(drug_df[ws_col], errors="coerce").dropna()
            if len(drug_numeric):
                stats["per_drug_stats"][str(drug)] = {
                    "count": int(len(drug_df)),
                    "avg_score": float(round(drug_numeric.mean(), 2)),
                    "min_score": float(round(drug_numeric.min(), 2)),
                    "max_score": float(round(drug_numeric.max(), 2)),
                }

    for sf_col in ["SF1 Score", "SF2 Score", "SF3 Score", "SF4 Score"]:
        if sf_col in df.columns:
            sf_numeric = pd.to_numeric(df[sf_col], errors="coerce").dropna()
            if len(sf_numeric):
                stats[f"avg_{sf_col.lower().replace(' ', '_')}"] = float(round(sf_numeric.mean(), 2))

    return stats


# ── Country score statistics ──────────────────────────────────────────────────

def compute_country_statistics(df_country: pd.DataFrame) -> dict:
    """Derive summary data from patent_strength_country_score_table."""
    result = {
        "by_drug": {},
        "final_scores": {},
        "all_jurisdictions": [],
    }
    if df_country.empty:
        return result

    result["all_jurisdictions"] = sorted(df_country["Jurisdiction"].dropna().unique().tolist())

    for drug, grp in df_country.groupby("Drug Name"):
        jur_map = {}
        final_score = None
        for _, row in grp.iterrows():
            jur = str(row.get("Jurisdiction", ""))
            jur_map[jur] = {
                "country_name":           str(row.get("Country Name", "")),
                "country_weight":         row.get("Country Weight"),
                "patent_count":           row.get("# Patents"),
                "avg_weighted_score":     row.get("Avg Weighted Score"),
                "country_weighted_score": row.get("Country Weighted Score"),
                "final_patent_score":     row.get("Final Patent Score (Drug Total)"),
            }
            if final_score is None:
                final_score = row.get("Final Patent Score (Drug Total)")
        result["by_drug"][str(drug)] = jur_map
        result["final_scores"][str(drug)] = final_score

    return result


# ── LLM narrative generation ──────────────────────────────────────────────────

def generate_executive_summary(stats: dict, df_final: pd.DataFrame, country_stats: dict = None) -> dict:
    """Use Gemini to produce a detailed executive summary and findings narrative."""
    patent_rows = []
    for _, r in df_final.iterrows():
        patent_rows.append({
            "drug": str(r.get("Drug Name", "")),
            "patent": str(r.get("Patent Number", "")),
            "type": str(r.get("Patent Type", "")),
            "jurisdiction": str(r.get("Jurisdiction", "N/A")),
            "sf1": r.get("SF1 Score", "N/A"),
            "sf2": r.get("SF2 Score", "N/A"),
            "sf3": r.get("SF3 Score", "N/A"),
            "sf4": r.get("SF4 Score", "N/A"),
            "weighted": r.get("Weighted Final Score", "N/A"),
            "core_step": str(r.get("Core Inventive Step", ""))[:300],
            "vulnerabilities": str(r.get("Key Vulnerabilities", ""))[:300],
            "strengths": str(r.get("Key Strengths", ""))[:300],
        })

    per_drug_json   = json.dumps(stats.get('per_drug_stats') or {})
    score_dist_json = json.dumps(stats.get('score_distribution') or {})

    hsp = stats.get("highest_score_patent") or {}
    hsp_summary = (
        f"Highest weighted score patent: {hsp.get('patent_number', 'N/A')} "
        f"(Drug: {hsp.get('drug', 'N/A')}, Jurisdiction: {hsp.get('jurisdiction', 'N/A')}, "
        f"Score: {hsp.get('weighted_score', 'N/A')}). "
        f"Vulnerabilities: {hsp.get('vulnerabilities', 'N/A')[:300]}"
        if hsp else "N/A"
    )

    hspj = stats.get("highest_score_per_jurisdiction") or {}
    hspj_summary = "; ".join(
        f"{jur}: {v.get('patent_number', 'N/A')} (Score: {v.get('weighted_score', 'N/A')})"
        for jur, v in hspj.items()
    ) if hspj else "N/A"

    country_context = "N/A"
    if country_stats and country_stats.get("by_drug"):
        lines = []
        for drug, jur_map in country_stats["by_drug"].items():
            final = country_stats["final_scores"].get(drug, "N/A")
            lines.append(f"Drug: {drug}  |  Final Patent Score: {final}")
            for jur, jdata in sorted(jur_map.items(), key=lambda x: -(x[1].get("country_weight") or 0)):
                lines.append(
                    f"  {jur} ({jdata['country_name']}): weight={jdata['country_weight']}, "
                    f"patents={jdata['patent_count']}, avg_weighted={jdata['avg_weighted_score']}, "
                    f"country_weighted={jdata['country_weighted_score']}"
                )
        country_context = "\n".join(lines)

    prompt = f"""You are a senior pharmaceutical patent analyst writing a detailed report.
Based on the data below, produce a thorough analytical report. Be specific — reference
patent numbers, drug names, jurisdictions, and actual scores throughout your analysis.

PORTFOLIO STATISTICS:
- Total blocking patents analysed: {stats['total_patents']}
- Drugs covered: {', '.join(stats['drugs'])}
- Average weighted robustness score: {stats['avg_weighted']} (1 = Very Robust, 5 = Highly Vulnerable)
- Score distribution: {score_dist_json}
- Per-drug averages: {per_drug_json}
- Sub-factor averages: SF1 (Novelty)={stats.get('avg_sf1_score', 'N/A')}, SF2 (Obvious-to-Combine)={stats.get('avg_sf2_score', 'N/A')}, SF3 (Prosecution History)={stats.get('avg_sf3_score', 'N/A')}, SF4 (Secondary Considerations)={stats.get('avg_sf4_score', 'N/A')}
- {hsp_summary}
- Highest score patent per jurisdiction: {hspj_summary}

COUNTRY-WEIGHTED SCORES (from patent_strength_country_score_table):
{country_context}

PATENT-LEVEL DATA (JSON):
{json.dumps(patent_rows, indent=1)}

Respond ONLY with a valid JSON object (no markdown fences):
{{
  "executive_summary": "<5-8 sentences providing a thorough overview of the portfolio's legal robustness, mentioning specific drugs and overall risk posture>",
  "key_findings": ["<detailed finding 1 with patent numbers>", "<detailed finding 2>", "<detailed finding 3>", "<detailed finding 4>", "<detailed finding 5>", "<detailed finding 6>"],
  "risk_highlights": "<4-6 sentences detailing the highest-risk patents, their specific vulnerabilities, and why they are exposed to challenge. Reference patent numbers and scores.>",
  "strength_highlights": "<4-6 sentences detailing the most robust patents, their defensive strengths, and what makes them resilient. Reference patent numbers and scores.>",
  "sf_analysis": "<4-6 sentences analysing sub-factor trends across the portfolio. Which sub-factor is weakest/strongest overall? Are there patterns by drug or patent type?>",
  "highest_score_narrative": "<4-6 sentences providing an in-depth analysis of the patent with the highest weighted score: why it scores so high, its core inventive weaknesses, the specific legal risks it faces, and strategic implications for challengers or defenders.>",
  "country_score_narrative": "<3-5 sentences analysing the country-weighted final patent scores per drug: which jurisdictions drive the highest risk, how country weights affect the overall score, and strategic geographic implications>",
  "per_drug_narratives": {{
    "<drug_name>": "<3-5 sentence analysis of this drug's patent portfolio robustness, key risks, and protective strengths>"
  }}
}}
"""
    text = call_gemini(prompt)
    result = _extract_json(text)
    if result:
        return result
    return {
        "executive_summary": "Analysis complete. See tables below for details.",
        "key_findings": ["See detailed tables."],
        "risk_highlights": "Refer to score tables.",
        "strength_highlights": "Refer to score tables.",
        "sf_analysis": "See sub-factor scores in the table.",
        "highest_score_narrative": "See score tables for details on the highest-scoring patent.",
        "country_score_narrative": "See country-weighted score table for details.",
        "per_drug_narratives": {},
    }


def apply_rationale_column(df: pd.DataFrame, per_drug_narratives: dict) -> pd.DataFrame:
    """
    Stamp each row of df with the Drug-Level Analysis narrative for its drug.

    The narrative written under Drug-Level Analysis for a given drug is inserted
    into a new 'rationale' column for every row belonging to that drug.
    """
    df = df.copy()
    df["rationale"] = ""
    for drug_name, narrative_text in per_drug_narratives.items():
        mask = df["Drug Name"] == drug_name
        df.loc[mask, "rationale"] = narrative_text
    return df


def ensure_rationale_column(client: bigquery.Client) -> None:
    """
    Add a STRING 'rationale' column to patent_strength_table if it doesn't
    already exist. BigQuery ALTER TABLE ADD COLUMN is idempotent-safe when
    the column is absent, but raises an error if it already exists, so we
    check the schema first.
    """
    table_ref = f"{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_STRENGTH_TABLE}"
    table = client.get_table(table_ref)
    existing_cols = {field.name for field in table.schema}

    if "rationale" in existing_cols:
        print("  'rationale' column already exists in BigQuery table — skipping ALTER.")
        return

    print("  Adding 'rationale' column to BigQuery table...")
    alter_sql = f"ALTER TABLE `{table_ref}` ADD COLUMN rationale STRING"
    client.query(alter_sql).result()
    print("  Column added successfully.")


def write_rationale_to_bigquery(per_drug_narratives: dict) -> None:
    """
    Persist the drug-level rationale text back to patent_strength_table in
    BigQuery.  One parameterised UPDATE is issued per drug so only the rows
    for that drug are touched.

    Steps:
      1. Ensure the 'rationale' column exists (ALTER TABLE if needed).
      2. Run a parameterised UPDATE … WHERE drug_name = @drug_name for each drug.
    """
    if not per_drug_narratives:
        print("  No per-drug narratives to write — skipping BigQuery update.")
        return

    client    = _bq_client()
    table_ref = f"{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_STRENGTH_TABLE}"

    # Step 1 — guarantee the column exists
    ensure_rationale_column(client)

    # Step 2 — update each drug's rows
    print(f"  Writing rationale to BigQuery for {len(per_drug_narratives)} drug(s)...")
    for drug_name, narrative_text in per_drug_narratives.items():
        sql = f"""
            UPDATE `{table_ref}`
            SET rationale = @rationale
            WHERE drug_name = @drug_name
        """
        job_config = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("rationale",  "STRING", narrative_text),
                bigquery.ScalarQueryParameter("drug_name",  "STRING", drug_name),
            ]
        )
        try:
            client.query(sql, job_config=job_config).result()
            print(f"    ✅ rationale written for drug: {drug_name}")
        except Exception as e:
            print(f"    [ERROR] Failed to update rationale for drug '{drug_name}': {e}")
            raise

    print("  BigQuery rationale update complete.")


def ensure_timestamp_columns(client: bigquery.Client) -> None:
    """
    Add TIMESTAMP columns 'created_at' and 'updated_at' to patent_strength_table
    if they don't already exist.

    - created_at: set only on the first run (when the column value is NULL).
    - updated_at: overwritten on every subsequent run.
    """
    table_ref = f"{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_STRENGTH_TABLE}"
    table = client.get_table(table_ref)
    existing_cols = {field.name for field in table.schema}

    for col in ("created_at", "updated_at"):
        if col in existing_cols:
            print(f"  '{col}' column already exists in BigQuery table — skipping ALTER.")
        else:
            print(f"  Adding '{col}' column to BigQuery table...")
            alter_sql = f"ALTER TABLE `{table_ref}` ADD COLUMN {col} TIMESTAMP"
            client.query(alter_sql).result()
            print(f"  '{col}' column added successfully.")


def write_timestamps_to_bigquery(drug_names: list, ts: datetime = None) -> None:
    """
    Write created_at and updated_at timestamps to patent_strength_table.

    - created_at: written only when the existing value is NULL (i.e. first run).
                  Existing values are preserved on subsequent runs.
    - updated_at: always overwritten with the current UTC time on every run.

    Steps:
      1. Ensure both timestamp columns exist (ALTER TABLE if needed).
      2. Run parameterised UPDATEs per drug.
    """
    if not drug_names:
        print("  No drug names provided — skipping timestamp update.")
        return

    if ts is None:
        ts = datetime.now(tz=timezone.utc)

    client    = _bq_client()
    table_ref = f"{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_STRENGTH_TABLE}"

    ensure_timestamp_columns(client)

    print(f"  Writing timestamps ({ts.isoformat()}) to BigQuery for {len(drug_names)} drug(s)...")
    for drug_name in drug_names:
        # created_at — only set when currently NULL (preserves the original first-run value)
        sql_created = f"""
            UPDATE `{table_ref}`
            SET created_at = @ts
            WHERE drug_name = @drug_name
              AND created_at IS NULL
        """
        # updated_at — always written to reflect the current run time
        sql_updated = f"""
            UPDATE `{table_ref}`
            SET updated_at = @ts
            WHERE drug_name = @drug_name
        """
        job_cfg = bigquery.QueryJobConfig(
            query_parameters=[
                bigquery.ScalarQueryParameter("ts",        "TIMESTAMP", ts),
                bigquery.ScalarQueryParameter("drug_name", "STRING",    drug_name),
            ]
        )
        try:
            client.query(sql_created, job_config=job_cfg).result()
            client.query(sql_updated, job_config=job_cfg).result()
            print(f"    ✅ timestamps written for drug: {drug_name}")
        except Exception as e:
            print(f"    [ERROR] Failed to update timestamps for drug '{drug_name}': {e}")
            raise

    print("  BigQuery timestamp update complete.")


# ── Document builder ──────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_el = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color,
        qn("w:val"): "clear",
    })
    shading.append(shading_el)


def _build_hsp_table(doc, hsp: dict, shade_color: str = "F9EBEB"):
    """Build and return a two-column info table for a highest-score patent dict."""
    try:
        score_label = SCORE_LABEL.get(int(round(float(hsp.get("weighted_score", 0)))), "N/A")
    except (ValueError, TypeError):
        score_label = "N/A"

    hsp_items = [
        ("Patent Number",                       hsp.get("patent_number", "N/A")),
        ("Drug Name",                            hsp.get("drug", "N/A")),
        ("Patent Type",                          hsp.get("patent_type", "N/A")),
        ("Jurisdiction",                         hsp.get("jurisdiction", "N/A")),
        ("Weighted Final Score",
         f"{hsp.get('weighted_score', 'N/A')} / 5.0  —  {score_label}"),
        ("SF1 Score (Novelty)",                  str(hsp.get("sf1", "N/A"))),
        ("SF2 Score (Obvious-to-Combine)",       str(hsp.get("sf2", "N/A"))),
        ("SF3 Score (Prosecution History)",      str(hsp.get("sf3", "N/A"))),
        ("SF4 Score (Secondary Considerations)", str(hsp.get("sf4", "N/A"))),
    ]

    tbl = doc.add_table(rows=len(hsp_items), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for i, (label, value) in enumerate(hsp_items):
        cell_l = tbl.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell_l, shade_color)

        cell_r = tbl.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        if label == "Weighted Final Score":
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    for row_obj in tbl.rows:
        row_obj.cells[0].width = Inches(2.5)
        row_obj.cells[1].width = Inches(4.2)

    return tbl


def build_report(data: dict, output_path: str):
    """Build the Word document."""
    df_final = data["final"]
    if df_final.empty:
        print("ERROR: patent_strength_table is empty or missing. Cannot generate report.")
        return

    df_country    = data.get("country_scores", pd.DataFrame())
    stats         = compute_statistics(df_final)
    country_stats = compute_country_statistics(df_country)

    print("Generating narrative with Gemini 2.5 Flash...")
    narrative = generate_executive_summary(stats, df_final, country_stats)

    # ── Stamp rationale column onto df_final ──────────────────────────────────
    per_drug_narratives = narrative.get("per_drug_narratives", {})
    df_final = apply_rationale_column(df_final, per_drug_narratives)
    # Propagate the updated df back into data so callers can use it
    data["final"] = df_final

    print("Writing rationale back to BigQuery...")
    write_rationale_to_bigquery(per_drug_narratives)

    print("Writing timestamps (created_at / updated_at) back to BigQuery...")
    drug_names = df_final["Drug Name"].dropna().unique().tolist() if not df_final.empty else []
    write_timestamps_to_bigquery(drug_names)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.6)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font  = style.font
    font.name  = "Arial"
    font.size  = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after  = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated {datetime.now().strftime('%B %d, %Y')}  •  "
        f"{stats['total_patents']} Blocking Patents  •  "
        f"{len(stats['drugs'])} Drug(s)"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.space_after = Pt(10)

    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(6)
    pBdr   = p_line._element.get_or_add_pPr().makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6",
        qn("w:space"): "1", qn("w:color"): "1F3864",
    })
    pBdr.append(bottom)
    p_line._element.get_or_add_pPr().append(pBdr)

    # ── Executive Summary ─────────────────────────────────────────────────────
    h = doc.add_heading("Executive Summary", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)
    p = doc.add_paragraph(narrative.get("executive_summary", ""))
    p.paragraph_format.space_after = Pt(6)

    # ── Portfolio Overview ────────────────────────────────────────────────────
    h = doc.add_heading("Portfolio Overview", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)

    overview_items = [
        ("Total Patents Analysed", str(stats["total_patents"])),
        ("Drugs Covered",          ", ".join(stats["drugs"])),
        ("Average Weighted Score", f"{stats['avg_weighted']} / 5.0" if stats["avg_weighted"] else "N/A"),
    ]
    sf_labels = [
        ("avg_sf1_score", "Avg SF1 (Novelty)"),
        ("avg_sf2_score", "Avg SF2 (Obvious-to-Combine)"),
        ("avg_sf3_score", "Avg SF3 (Prosecution History)"),
        ("avg_sf4_score", "Avg SF4 (Secondary Considerations)"),
    ]
    for key, label in sf_labels:
        val = stats.get(key)
        if val is not None:
            overview_items.append((label, f"{val} / 5.0"))

    dist_parts = []
    for s in range(1, 6):
        count = stats["score_distribution"].get(s, 0)
        if count > 0:
            dist_parts.append(f"{SCORE_LABEL[s]}: {count}")
    if dist_parts:
        overview_items.append(("Score Distribution", "; ".join(dist_parts)))

    overview_table = doc.add_table(rows=len(overview_items), cols=2)
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview_table.style     = "Table Grid"
    for i, (label, value) in enumerate(overview_items):
        cell_l = overview_table.rows[i].cells[0]
        cell_l.text = ""
        p = cell_l.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell_l, "E8EDF3")

        cell_r = overview_table.rows[i].cells[1]
        cell_r.text = ""
        p = cell_r.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)

    for row_obj in overview_table.rows:
        row_obj.cells[0].width = Inches(2.5)
        row_obj.cells[1].width = Inches(4.2)

    doc.add_paragraph("")

    # ── Key Findings ──────────────────────────────────────────────────────────
    h = doc.add_heading("Key Findings", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)
    for finding in narrative.get("key_findings", []):
        p = doc.add_paragraph(finding, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # ── Patent Scores Table (now includes Rationale column) ───────────────────
    h = doc.add_heading("Patent Score Summary", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)

    cols         = ["Drug Name", "Patent Number", "SF1 Score", "SF2 Score",
                    "SF3 Score", "SF4 Score", "Weighted Final Score", "Created At", "Updated At", "rationale"]
    display_cols = ["Drug", "Patent No.", "SF1", "SF2", "SF3", "SF4", "Final", "Created At", "Updated At", "Rationale"]

    df_table = df_final.copy()
    ws_col   = "Weighted Final Score"
    if ws_col in df_table.columns:
        df_table["_ws_numeric"] = pd.to_numeric(df_table[ws_col], errors="coerce")
        df_table = df_table.dropna(subset=["_ws_numeric"])
        df_table = df_table.drop(columns=["_ws_numeric"])

    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    for i, label in enumerate(display_cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    for _, row in df_table.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            val  = row.get(col, "")
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            # Rationale is left-aligned; all others centred
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col == "rationale" else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if pd.notna(val) else "N/A")
            run.font.size = Pt(8)
            if i >= 2 and col not in ("rationale", "Created At", "Updated At"):
                try:
                    score_int = int(round(float(val)))
                    if score_int in SCORE_COLOR_MAP:
                        run.font.color.rgb = SCORE_COLOR_MAP[score_int]
                        run.bold = True
                except (ValueError, TypeError):
                    pass

    # Column widths — rationale gets generous space; others stay narrow
    widths = [Inches(1.0), Inches(1.0), Inches(0.45), Inches(0.45),
              Inches(0.45), Inches(0.45), Inches(0.55), Inches(0.8), Inches(0.8), Inches(1.65)]
    for row_obj in table.rows:
        for i, cell in enumerate(row_obj.cells):
            cell.width = widths[i]

    doc.add_paragraph("")

    # ── Highest Weighted Score Patent (Overall) ───────────────────────────────
    hsp = stats.get("highest_score_patent")
    if hsp:
        h = doc.add_heading("Highest Weighted Score Patent", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.size = Pt(13)

        _build_hsp_table(doc, hsp, shade_color="F9EBEB")
        doc.add_paragraph("")

        if hsp.get("vulnerabilities") and hsp["vulnerabilities"] != "N/A":
            p = doc.add_paragraph()
            run = p.add_run("Key Vulnerabilities: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.add_run(hsp["vulnerabilities"]).font.size = Pt(10)

        if hsp.get("strengths") and hsp["strengths"] != "N/A":
            p = doc.add_paragraph()
            run = p.add_run("Key Strengths: ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            p.add_run(hsp["strengths"]).font.size = Pt(10)

        if hsp.get("core_step") and hsp["core_step"] != "N/A":
            p = doc.add_paragraph()
            run = p.add_run("Core Inventive Step: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(hsp["core_step"]).font.size = Pt(10)

        doc.add_paragraph("")

        highest_narrative = narrative.get("highest_score_narrative", "")
        if highest_narrative:
            p = doc.add_paragraph()
            run = p.add_run("In-Depth Analysis")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            doc.add_paragraph(highest_narrative)

        doc.add_paragraph("")

    # ── Highest Weighted Score Patent per Jurisdiction ────────────────────────
    hspj = stats.get("highest_score_per_jurisdiction") or {}
    if hspj:
        h = doc.add_heading("Highest Weighted Score Patent by Jurisdiction", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.size = Pt(13)

        p_intro = doc.add_paragraph(
            "The table below identifies the most legally vulnerable patent within each "
            "jurisdiction across the analysed portfolio."
        )
        p_intro.paragraph_format.space_after = Pt(6)

        for jur, jur_hsp in hspj.items():
            h3 = doc.add_heading(f"Jurisdiction: {jur}", level=3)
            for run in h3.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                run.font.size = Pt(11)
            _build_hsp_table(doc, jur_hsp, shade_color="F9EBEB")
            doc.add_paragraph("")

    # ── Country-Weighted Score Section ────────────────────────────────────────
    if country_stats and country_stats.get("by_drug"):
        h = doc.add_heading("Country-Weighted Patent Scores", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            run.font.size = Pt(13)

        intro_p = doc.add_paragraph(
            "The table below shows jurisdiction-level weighted scores from "
            "patent_strength_country_score_table. Country weights reflect the "
            "strategic commercial importance of each jurisdiction."
        )
        intro_p.paragraph_format.space_after = Pt(6)

        for drug, jur_map in country_stats["by_drug"].items():
            h3 = doc.add_heading(drug, level=3)
            for run in h3.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                run.font.size = Pt(11)

            final_score = country_stats["final_scores"].get(drug)
            if final_score is not None:
                fs_p = doc.add_paragraph()
                fs_run = fs_p.add_run(f"Final Patent Score (Drug Total): {round(float(final_score), 4)}")
                fs_run.bold = True
                fs_run.font.size = Pt(10)
                fs_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

            ct_headers = ["Jurisdiction", "Country", "Weight", "# Patents",
                          "Avg Weighted Score", "Country Weighted Score"]
            ct_widths  = [Inches(0.8), Inches(1.5), Inches(0.6), Inches(0.6),
                          Inches(1.3), Inches(1.5)]

            ct = doc.add_table(rows=1, cols=len(ct_headers))
            ct.alignment = WD_TABLE_ALIGNMENT.CENTER
            ct.style     = "Table Grid"
            for i, label in enumerate(ct_headers):
                cell = ct.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "1F3864")

            for jur, jdata in sorted(jur_map.items(),
                                     key=lambda x: -(x[1].get("country_weight") or 0)):
                row_cells = ct.add_row().cells
                vals = [
                    jur,
                    str(jdata.get("country_name", "")),
                    str(jdata.get("country_weight", "")),
                    str(jdata.get("patent_count", "")),
                    str(jdata.get("avg_weighted_score", "N/A")),
                    str(round(float(jdata["country_weighted_score"]), 4))
                    if jdata.get("country_weighted_score") is not None else "N/A",
                ]
                for i, val in enumerate(vals):
                    cell = row_cells[i]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(val)
                    run.font.size = Pt(8)
                    if i == 4:
                        try:
                            s_int = int(round(float(val)))
                            if s_int in SCORE_COLOR_MAP:
                                run.font.color.rgb = SCORE_COLOR_MAP[s_int]
                                run.bold = True
                        except (ValueError, TypeError):
                            pass

            for row_obj in ct.rows:
                for i, cell in enumerate(row_obj.cells):
                    cell.width = ct_widths[i]

            doc.add_paragraph("")

        csn = narrative.get("country_score_narrative", "")
        if csn:
            p = doc.add_paragraph()
            run = p.add_run("Geographic Risk Analysis")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            doc.add_paragraph(csn)

        doc.add_paragraph("")

    # ── Sub-Factor Analysis ───────────────────────────────────────────────────
    sf_text = narrative.get("sf_analysis", "")
    if sf_text:
        h = doc.add_heading("Sub-Factor Analysis", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            run.font.size = Pt(13)
        doc.add_paragraph(sf_text)

    # ── Risk & Strength Highlights ────────────────────────────────────────────
    h = doc.add_heading("Risk & Strength Analysis", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)

    p = doc.add_paragraph()
    run = p.add_run("Highest Risk Patents")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    doc.add_paragraph(narrative.get("risk_highlights", "N/A"))

    p = doc.add_paragraph()
    run = p.add_run("Most Robust Patents")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    doc.add_paragraph(narrative.get("strength_highlights", "N/A"))

    # ── Per-Drug Breakdown ────────────────────────────────────────────────────
    per_drug = narrative.get("per_drug_narratives", {})
    if per_drug:
        h = doc.add_heading("Drug-Level Analysis", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            run.font.size = Pt(13)

        for drug_name, drug_narrative in per_drug.items():
            h3 = doc.add_heading(drug_name, level=3)
            for run in h3.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                run.font.size = Pt(11)

            drug_stat = stats.get("per_drug_stats", {}).get(drug_name, {})
            if drug_stat:
                stat_p = doc.add_paragraph()
                stat_run = stat_p.add_run(
                    f"Patents: {drug_stat.get('count', 'N/A')}  |  "
                    f"Avg Score: {drug_stat.get('avg_score', 'N/A')}  |  "
                    f"Range: {drug_stat.get('min_score', 'N/A')} – {drug_stat.get('max_score', 'N/A')}"
                )
                stat_run.font.size = Pt(9)
                stat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                stat_run.italic = True

            doc.add_paragraph(drug_narrative)

    # ── Sub-Factor Weights Table ──────────────────────────────────────────────
    h = doc.add_heading("Sub-Factor Scoring Framework", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)

    sf_framework = [
        ("1", "Novelty & Non-Obviousness",
         "Closeness of the claimed invention to prior art structurally or conceptually.",
         "40%"),
        ("2", "Obvious-to-Combine Risk",
         "Likelihood a skilled person would combine known elements with a reasonable expectation of success.",
         "30%"),
        ("3", "Prosecution History Vulnerability",
         "Extent of claim narrowing or limiting arguments during prosecution that may weaken enforceability.",
         "20%"),
        ("4", "Secondary Considerations",
         "Objective evidence supporting non-obviousness (e.g. commercial success, long-felt need).",
         "10%"),
    ]

    sf_table    = doc.add_table(rows=1, cols=4)
    sf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sf_table.style     = "Table Grid"
    sf_headers         = ["SF #", "Name", "Description", "Weight"]
    sf_col_widths      = [Inches(0.45), Inches(1.6), Inches(3.7), Inches(0.65)]

    for i, label in enumerate(sf_headers):
        cell = sf_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    for sf_num, sf_name, sf_desc, sf_weight in sf_framework:
        row_cells = sf_table.add_row().cells

        row_cells[0].text = ""
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sf_num)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(row_cells[0], "E8EDF3")

        row_cells[1].text = ""
        p = row_cells[1].paragraphs[0]
        run = p.add_run(sf_name)
        run.bold = True
        run.font.size = Pt(9)

        row_cells[2].text = ""
        p = row_cells[2].paragraphs[0]
        p.add_run(sf_desc).font.size = Pt(9)

        row_cells[3].text = ""
        p = row_cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sf_weight)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    for row_obj in sf_table.rows:
        for i, cell in enumerate(row_obj.cells):
            cell.width = sf_col_widths[i]

    doc.add_paragraph("")

    # ── Score Legend ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("Score Legend: ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    legend_text = "  |  ".join(f"{k} = {v}" for k, v in SCORE_LABEL.items())
    run = p.add_run(legend_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Footer note ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(
        "This report was auto-generated from Patent Legal Robustness Scorer output "
        "using Gemini 2.5 Flash for narrative analysis."
    )
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Save locally ──────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"\n✅ Report saved locally → {output_path}")


# ── GCS Upload ────────────────────────────────────────────────────────────────

def upload_to_gcs(local_path: str, drug_names: list) -> list:
    """
    Upload the generated .docx to GCS under each drug's folder.

    Destination path per drug:
        gs://cognito-gcs/Cognito_new/reports/{drug_name}/Patent_Strength_Analysis.docx

    Uses service account file or ADC (Cloud Run).
    Requires: pip install google-cloud-storage
    """
    try:
        from google.cloud import storage
    except ImportError:
        raise ImportError(
            "google-cloud-storage is required.\n"
            "Run: pip install google-cloud-storage"
        )

    credentials = _get_credentials()
    client   = storage.Client(project=BQ_PROJECT_ID, credentials=credentials)
    bucket   = client.bucket(GCS_BUCKET)
    gcs_uris = []

    for drug_name in drug_names:
        safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", str(drug_name))
        blob_name = f"{GCS_BASE_PATH}/{safe_name}/{GCS_FILE_NAME}"
        gcs_uri   = f"gs://{GCS_BUCKET}/{blob_name}"

        print(f"  Uploading to GCS: {gcs_uri}")
        try:
            blob = bucket.blob(blob_name)
            blob.upload_from_filename(
                local_path,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            print(f"  Upload successful: {gcs_uri}")
            gcs_uris.append(gcs_uri)
        except Exception as e:
            print(f"  [ERROR] GCS upload failed for drug '{drug_name}': {e}")
            raise

    return gcs_uris


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a detailed Word report from BigQuery patent scorer tables"
    )
    parser.add_argument("--output", "-o", default="patent_robustness_report.docx",
                        help="Output .docx path (default: patent_robustness_report.docx)")
    args = parser.parse_args()

    if not API_KEY:
        raise EnvironmentError(
            "GEMINI_API_KEY is not set. Add it to your .env file:\n"
            "  GEMINI_API_KEY=your-key"
        )
    data = load_from_bigquery()
    build_report(data, args.output)

    # ── Upload to GCS ─────────────────────────────────────────────────────────
    df_final   = data.get("final", pd.DataFrame())
    drug_names = df_final["Drug Name"].dropna().unique().tolist() if not df_final.empty else []

    if drug_names:
        print(f"\nUploading report to GCS for {len(drug_names)} drug(s)...")
        gcs_uris = upload_to_gcs(args.output, drug_names)

        print(f"\n{'='*60}")
        print("  GCS Upload Summary:")
        for uri in gcs_uris:
            print(f"    {uri}")
        print("=" * 60)
    else:
        print("\n[WARN] No drug names found — skipping GCS upload.")


if __name__ == "__main__":
    main()
